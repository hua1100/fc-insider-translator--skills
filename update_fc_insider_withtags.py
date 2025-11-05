#!/usr/bin/env python3
"""
FC Insider 翻译更新脚本（正确处理 Tag 样式）
- 跳过 Tag 样式的 runs
- 只读取实际的翻译内容
- 支持追踪修订
"""

import json
import argparse
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要安装 python-docx")
    print("运行: pip install --user python-docx")
    exit(1)

from tag_protector import protect_tags, restore_tags


def get_cell_text_without_tags(cell):
    """
    获取单元格文本，跳过 Tag 样式的 runs

    Args:
        cell: Word 表格单元格

    Returns:
        不包含 Tag 样式的文本内容
    """
    text_parts = []

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            # 检查是否有 rStyle
            run_element = run._element
            rpr = run_element.find(qn('w:rPr'))

            if rpr is not None:
                r_style = rpr.find(qn('w:rStyle'))
                if r_style is not None:
                    style_val = r_style.get(qn('w:val'))
                    # 跳过 Tag 样式
                    if style_val == 'Tag':
                        continue

            # 添加非 Tag 样式的文本
            if run.text:
                text_parts.append(run.text)

    return ''.join(text_parts).strip()


def enable_track_changes(doc):
    """启用文档的追踪修订功能"""
    settings = doc.settings
    settings_element = settings.element

    track_revisions = settings_element.find(qn('w:trackRevisions'))
    if track_revisions is None:
        track_revisions = OxmlElement('w:trackRevisions')
        settings_element.append(track_revisions)
        print("✓ 文档层级追踪修订已启用")
    else:
        print("✓ 文档层级追踪修订已存在")


def create_delete_element(author, date, text_content, revision_id, rpr_element=None):
    """创建删除标记元素"""
    del_element = OxmlElement('w:del')
    del_element.set(qn('w:id'), str(revision_id))
    del_element.set(qn('w:author'), author)
    del_element.set(qn('w:date'), date)

    del_run = OxmlElement('w:r')

    if rpr_element is not None:
        # 复制格式，但移除 rStyle（避免复制 Tag 样式）
        import copy
        rpr_copy = copy.deepcopy(rpr_element)
        # 尝试移除 rStyle
        for r_style in rpr_copy.findall(qn('w:rStyle')):
            if r_style.get(qn('w:val')) == 'Tag':
                rpr_copy.remove(r_style)
        del_run.append(rpr_copy)

    del_text = OxmlElement('w:delText')
    del_text.text = text_content
    del_text.set(qn('xml:space'), 'preserve')

    del_run.append(del_text)
    del_element.append(del_run)

    return del_element


def create_insert_element(author, date, text_content, revision_id, rpr_element=None):
    """创建插入标记元素"""
    ins_element = OxmlElement('w:ins')
    ins_element.set(qn('w:id'), str(revision_id))
    ins_element.set(qn('w:author'), author)
    ins_element.set(qn('w:date'), date)

    ins_run = OxmlElement('w:r')

    if rpr_element is not None:
        # 复制格式，但移除 Tag 样式
        import copy
        rpr_copy = copy.deepcopy(rpr_element)
        for r_style in rpr_copy.findall(qn('w:rStyle')):
            if r_style.get(qn('w:val')) == 'Tag':
                rpr_copy.remove(r_style)
        ins_run.append(rpr_copy)

    ins_text = OxmlElement('w:t')
    ins_text.text = text_content
    ins_text.set(qn('xml:space'), 'preserve')

    ins_run.append(ins_text)
    ins_element.append(ins_run)

    return ins_element


def find_row_by_segment_id(table, segment_id):
    """在表格中查找包含指定 segment_id 的行"""
    for row_idx, row in enumerate(table.rows):
        # 同样跳过 Tag 样式读取 segment_id
        first_cell_text = get_cell_text_without_tags(row.cells[0])
        if first_cell_text == str(segment_id).strip():
            return row, row_idx
    return None, None


def replace_cell_with_track_changes(cell, old_text, new_text, author, date_str, revision_id):
    """
    使用追踪修订替换单元格内容（跳过 Tag 样式的内容）

    重要：只替换非 Tag 样式的内容，保留 Tag 样式的 runs
    """
    # 获取不包含 Tag 的当前文本
    current_text = get_cell_text_without_tags(cell)

    # 保护标签
    old_protected = protect_tags(old_text)
    new_protected = protect_tags(new_text)
    current_protected = protect_tags(current_text)

    # 验证文本（允许空 old_text）
    allow_empty_old = (old_text is None) or (str(old_text).strip() == "")

    if not allow_empty_old:
        # 简单的包含检查
        if old_protected.strip() not in current_protected and current_protected.strip() not in old_protected:
            raise ValueError(
                f"文本不匹配\n"
                f"  预期: '{old_text[:50]}...'\n"
                f"  实际: '{current_text[:50]}...'"
            )

    # 处理段落
    if len(cell.paragraphs) == 0:
        return revision_id

    # 找到第一个非 Tag 样式的段落
    target_paragraph = None
    for paragraph in cell.paragraphs:
        # 检查是否有非 Tag 内容
        has_non_tag_content = False
        for run in paragraph.runs:
            run_element = run._element
            rpr = run_element.find(qn('w:rPr'))
            if rpr is not None:
                r_style = rpr.find(qn('w:rStyle'))
                if r_style is not None and r_style.get(qn('w:val')) == 'Tag':
                    continue
            has_non_tag_content = True
            break

        if has_non_tag_content:
            target_paragraph = paragraph
            break

    if target_paragraph is None:
        # 如果没有找到，使用第一个段落
        target_paragraph = cell.paragraphs[0]

    p_element = target_paragraph._element

    # 保存格式（从第一个非 Tag run）
    rpr_element = None
    for run in target_paragraph.runs:
        run_element = run._element
        rpr = run_element.find(qn('w:rPr'))
        if rpr is not None:
            r_style = rpr.find(qn('w:rStyle'))
            if r_style is not None and r_style.get(qn('w:val')) == 'Tag':
                continue
            rpr_element = rpr
            break

    # 清空非 Tag 的内容，但保留 Tag runs
    for child in list(p_element):
        if child.tag == qn('w:r'):
            # 检查是否是 Tag 样式
            rpr = child.find(qn('w:rPr'))
            if rpr is not None:
                r_style = rpr.find(qn('w:rStyle'))
                if r_style is not None and r_style.get(qn('w:val')) == 'Tag':
                    # 保留 Tag 样式的 run
                    continue
            # 移除非 Tag 的 run
            p_element.remove(child)
        elif child.tag in [qn('w:del'), qn('w:ins'), qn('w:hyperlink')]:
            p_element.remove(child)

    # 添加删除标记
    if old_protected.strip():
        del_element = create_delete_element(
            author, date_str, old_protected, revision_id, rpr_element
        )
        p_element.append(del_element)
        revision_id += 1

    # 添加插入标记
    if new_protected.strip():
        ins_element = create_insert_element(
            author, date_str, new_protected, revision_id, rpr_element
        )
        p_element.append(ins_element)
        revision_id += 1

    return revision_id


def update_translations(doc, translations, author='Claude', verbose=False):
    """更新翻译"""
    successful = 0
    failed = []

    if not doc.tables:
        raise ValueError("文档中没有找到表格")

    table = doc.tables[0]
    date_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')
    revision_id = 0

    print(f"开始处理 {len(translations)} 个翻译...")
    print("=" * 80)

    for i, trans in enumerate(translations, 1):
        segment_id = trans.get('segment_id', 'Unknown')
        old_text = trans.get('old_text', '')
        new_text = trans.get('new_text', '')

        print(f"[{i}/{len(translations)}] 处理 {segment_id[:40]}...", end=' ')

        if verbose:
            print(f"\n  预期 old_text: '{old_text[:50]}...'")

        try:
            # 查找行
            row, row_idx = find_row_by_segment_id(table, segment_id)
            if row is None:
                raise ValueError(f"找不到 segment_id: {segment_id}")

            # 获取目标单元格
            if len(row.cells) < 4:
                raise ValueError("表格列数不足（需要至少 4 列）")

            target_cell = row.cells[3]

            # 诊断：显示实际内容
            if verbose:
                actual_text = get_cell_text_without_tags(target_cell)
                print(f"  实际文本（无 Tag）: '{actual_text[:50]}...'")

            # 应用追踪修订
            revision_id = replace_cell_with_track_changes(
                target_cell, old_text, new_text, author, date_str, revision_id
            )

            successful += 1
            print("✓")

        except Exception as e:
            error_msg = str(e)
            failed.append((segment_id, error_msg))

            if verbose:
                print(f"✗\n  详细错误: {error_msg}")
            else:
                short_error = error_msg.split('\n')[0][:60]
                print(f"✗ ({short_error})")

    return successful, failed


def restore_tags_in_document(doc):
    """恢复文档中的标签保护"""
    print("\n恢复标签保护...")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            run.text = restore_tags(run.text)

    print("✓ 标签恢复完成")


def main():
    parser = argparse.ArgumentParser(
        description='FC Insider 翻译更新（正确处理 Tag 样式）',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  python update_fc_insider_withtags.py --input input.docx --translations trans.json --output output.docx

  # 详细模式
  python update_fc_insider_withtags.py --input input.docx --translations trans.json --output output.docx --verbose
        '''
    )

    parser.add_argument('--input', required=True, help='输入 DOCX 文件')
    parser.add_argument('--translations', required=True, help='翻译 JSON 文件路径')
    parser.add_argument('--output', required=True, help='输出 DOCX 文件')
    parser.add_argument('--author', default='Claude', help='作者名称（默认: Claude）')
    parser.add_argument('--verbose', action='store_true', help='显示详细诊断信息')

    args = parser.parse_args()

    # 检查输入文件
    if not Path(args.input).exists():
        print(f"✗ 错误：输入文件不存在 - {args.input}")
        return 1

    if not Path(args.translations).exists():
        print(f"✗ 错误：翻译文件不存在 - {args.translations}")
        return 1

    try:
        # 加载翻译数据
        with open(args.translations, 'r', encoding='utf-8') as f:
            data = json.load(f)
            translations = data.get('translations', data)

        # 打开文档
        print(f"读取文档: {args.input}")
        doc = Document(args.input)

        # 启用追踪修订
        enable_track_changes(doc)

        print(f"FC Insider 翻译更新（正确处理 Tag 样式）")
        print(f"作者: {args.author}")
        print(f"翻译数量: {len(translations)}")
        print("=" * 80)

        # 更新翻译
        successful, failed = update_translations(doc, translations, args.author, args.verbose)

        # 恢复标签
        restore_tags_in_document(doc)

        print("=" * 80)
        print(f"✓ 更新完成: {successful}/{len(translations)}")

        if failed:
            print(f"✗ 失败: {len(failed)} 个")
            print("\n失敗詳情（前 10 個）:")
            for seg_id, error in failed[:10]:
                print(f"  • {seg_id[:40]}: {error[:200]}")

        print("=" * 80)

        # 保存文档
        print(f"\n保存文档: {args.output}")
        doc.save(args.output)
        print("✓ 文档已保存")

        return 0 if successful > 0 else 1

    except Exception as e:
        print(f"\n✗ 致命错误: {e}")
        import traceback
        traceback.print_exc()
        return 2


if __name__ == '__main__':
    import sys
    sys.exit(main())
