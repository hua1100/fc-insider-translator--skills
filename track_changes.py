#!/usr/bin/env python3

# -*- coding: utf-8 -*-

"""

應用Word追蹤修訂腳本

"""

from docx import Document

from docx.oxml import OxmlElement

from docx.oxml.ns import qn

from datetime import datetime

import json



def enable_track_changes(doc):

    """啟用文檔的追蹤修訂功能"""

    settings = doc.settings

    settings_element = settings.element

    

    # 檢查是否已有 trackRevisions 元素

    track_revisions = settings_element.find(qn('w:trackRevisions'))

    if track_revisions is None:

        track_revisions = OxmlElement('w:trackRevisions')

        settings_element.append(track_revisions)



def create_delete_element(author, date, text_content, revision_id):

    """創建刪除標記元素"""

    del_element = OxmlElement('w:del')

    del_element.set(qn('w:id'), str(revision_id))

    del_element.set(qn('w:author'), author)

    del_element.set(qn('w:date'), date)

    

    # 創建run元素

    del_run = OxmlElement('w:r')

    

    # 創建delText元素

    del_text = OxmlElement('w:delText')

    del_text.text = text_content

    del_text.set(qn('xml:space'), 'preserve')

    

    del_run.append(del_text)

    del_element.append(del_run)

    

    return del_element



def create_insert_element(author, date, text_content, revision_id):

    """創建插入標記元素"""

    ins_element = OxmlElement('w:ins')

    ins_element.set(qn('w:id'), str(revision_id))

    ins_element.set(qn('w:author'), author)

    ins_element.set(qn('w:date'), date)

    

    # 創建run元素

    ins_run = OxmlElement('w:r')

    

    # 創建text元素

    ins_text = OxmlElement('w:t')

    ins_text.text = text_content

    ins_text.set(qn('xml:space'), 'preserve')

    

    ins_run.append(ins_text)

    ins_element.append(ins_run)

    

    return ins_element



def replace_cell_with_track_changes(cell, new_text, author, date_str, revision_id):

    """使用追蹤修訂替換單元格內容"""

    # 獲取原始文字

    original_text = cell.text.strip()

    

    # 處理第一個段落

    if len(cell.paragraphs) == 0:

        return revision_id

    

    paragraph = cell.paragraphs[0]

    p_element = paragraph._element

    

    # 完全清空段落內容

    for child in list(p_element):

        if child.tag in [qn('w:r'), qn('w:del'), qn('w:ins'), 

                         qn('w:hyperlink'), qn('w:bookmarkStart'), 

                         qn('w:bookmarkEnd')]:

            p_element.remove(child)

    

    # 同時清空所有run的文字

    for run in paragraph.runs:

        run.text = ''

    

    # 只有在有原始文字時才添加刪除標記

    if original_text:

        del_element = create_delete_element(

            author, date_str, original_text, revision_id

        )

        p_element.append(del_element)

        revision_id += 1

    

    # 只有在有新文字時才添加插入標記

    if new_text.strip():

        ins_element = create_insert_element(

            author, date_str, new_text, revision_id

        )

        p_element.append(ins_element)

        revision_id += 1

    

    return revision_id



def apply_track_changes_to_translation(

    input_file_path,

    output_file_path,

    revised_translations,

    author="Translation Reviewer"

):

    """將翻譯修改應用到Word文檔，使用追蹤修訂功能"""

    

    # 讀取文檔

    doc = Document(input_file_path)

    

    # 啟用追蹤修訂

    enable_track_changes(doc)

    

    # 獲取表格

    if len(doc.tables) == 0:

        raise ValueError("文檔中沒有找到表格")

    

    table = doc.tables[0]

    

    # 獲取當前時間

    date_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')

    

    # 修訂ID計數器

    revision_id = 0

    

    # 處理每個修訂

    for revision in revised_translations:

        row_index = revision['row_index']

        new_text = revision['revised_text']

        

        # 確保row_index有效

        if row_index >= len(table.rows):

            print(f"警告：row_index {row_index} 超出範圍")

            continue

        

        # 獲取目標單元格（Target segment，第4列，索引3）

        row = table.rows[row_index]

        target_cell = row.cells[3]

        

        print(f"處理 Row {row_index}...")

        

        # 應用追蹤修訂

        revision_id = replace_cell_with_track_changes(

            target_cell,

            new_text,

            author,

            date_str,

            revision_id

        )

    

    # 保存文檔

    doc.save(output_file_path)

    

    print(f"\n✓ 已成功應用追蹤修訂到 {len(revised_translations)} 個片段")

    print(f"✓ 文件已保存到: {output_file_path}")

    print(f"✓ 使用的修訂ID: 0-{revision_id-1}")

    

    return output_file_path



# 主程序

if __name__ == "__main__":

    # 讀取修訂結果

    with open('revised_translations.json', 'r', encoding='utf-8') as f:

        revised_translations = json.load(f)

    

    # 應用追蹤修訂

    output_path = apply_track_changes_to_translation(

        input_file_path='/mnt/user-data/uploads/FCInsider_Dec2025_Issue9_2025Oct20_Englis_review.docx',

        output_file_path='/mnt/user-data/outputs/FCInsider_Dec2025_Issue9_修訂版.docx',

        revised_translations=revised_translations,

        author='Claude Translation Reviewer'

    )

    

    print(f"\n完成！修訂後的文件：{output_path}")