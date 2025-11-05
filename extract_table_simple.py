#!/usr/bin/env python3
"""
纯 Python 版本的 Word 表格提取器
不依赖 Pandoc，只使用 python-docx（可在 Claude Skills 环境中运行）

用法:
    python extract_table_simple.py input.docx output.md
    python extract_table_simple.py input.docx output.md --output-json table.json
"""

import sys
import json
import argparse
from pathlib import Path
from typing import List, Dict

try:
    from docx import Document
except ImportError:
    print("错误：需要安装 python-docx")
    print("运行: pip install --user python-docx")
    sys.exit(1)


def extract_table_to_markdown(docx_path: str, output_md: str) -> List[Dict[str, str]]:
    """
    从 Word 文档提取表格并转换为 Markdown

    Args:
        docx_path: Word 文档路径
        output_md: 输出 Markdown 文件路径

    Returns:
        提取的表格数据（列表形式）
    """
    print(f"读取 Word 文档: {docx_path}")
    doc = Document(docx_path)

    if not doc.tables:
        print("⚠️  警告：文档中没有找到表格")
        return []

    markdown_lines = ["# FC Insider Translation Table\n"]
    all_rows_data = []

    for table_idx, table in enumerate(doc.tables, 1):
        print(f"处理表格 {table_idx}...")
        markdown_lines.append(f"## Table {table_idx}\n")

        rows = table.rows
        if not rows:
            markdown_lines.append("*(空表格)*\n")
            continue

        # 提取表头（第一行）
        header_cells = rows[0].cells
        headers = [cell.text.strip() for cell in header_cells]

        # 生成 Markdown 表头
        if len(headers) >= 4:
            # FC Insider 标准四列格式
            markdown_lines.append("| Segment ID | Status | Source | Target |")
            markdown_lines.append("|------------|--------|--------|--------|")
        else:
            # 通用表格格式
            markdown_lines.append("| " + " | ".join(headers) + " |")
            markdown_lines.append("|" + "|".join(["---"] * len(headers)) + "|")

        # 提取数据行
        row_count = 0
        for row_idx, row in enumerate(rows[1:], 1):  # 跳过表头
            cells = row.cells
            cell_texts = [clean_cell_text(cell.text) for cell in cells]

            # 跳过空行
            if all(not text for text in cell_texts):
                continue

            # 确保至少有 4 列（FC Insider 格式）
            while len(cell_texts) < 4:
                cell_texts.append("")

            # 只取前 4 列
            cell_texts = cell_texts[:4]

            # 添加到 Markdown
            row_md = "| " + " | ".join(cell_texts) + " |"
            markdown_lines.append(row_md)

            # 保存为结构化数据
            if len(cell_texts) >= 4:
                all_rows_data.append({
                    'segment_id': cell_texts[0],
                    'status': cell_texts[1],
                    'source': cell_texts[2],
                    'target': cell_texts[3]
                })
                row_count += 1

        markdown_lines.append("")  # 表格之间空行
        print(f"  ✓ 提取 {row_count} 行数据")

    # 写入 Markdown 文件
    with open(output_md, 'w', encoding='utf-8') as f:
        f.write('\n'.join(markdown_lines))

    print(f"✓ Markdown 已保存: {output_md}")
    return all_rows_data


def clean_cell_text(text: str) -> str:
    """
    清理单元格文本
    - 移除多余空白
    - 将换行转为空格
    - 转义 Markdown 特殊字符
    """
    # 替换换行为空格
    text = text.replace('\n', ' ').replace('\r', ' ')

    # 移除多余空白
    text = ' '.join(text.split())

    # 转义 Markdown 管道符（避免破坏表格）
    text = text.replace('|', '\\|')

    return text


def main():
    parser = argparse.ArgumentParser(
        description='从 Word 文档提取表格为 Markdown（纯 Python 版本）',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  # 基本用法
  python extract_table_simple.py input.docx output.md

  # 同时生成 JSON
  python extract_table_simple.py input.docx output.md --output-json table.json
        '''
    )

    parser.add_argument('input_docx', help='输入 Word 文档路径')
    parser.add_argument('output_md', help='输出 Markdown 文件路径')
    parser.add_argument(
        '--output-json',
        help='可选：同时输出结构化 JSON 数据'
    )

    args = parser.parse_args()

    # 检查输入文件
    if not Path(args.input_docx).exists():
        print(f"✗ 错误：文件不存在 - {args.input_docx}")
        return 1

    print("=" * 80)
    print("Word 表格提取器（纯 Python 版本）")
    print("=" * 80)

    try:
        # 提取表格
        rows_data = extract_table_to_markdown(args.input_docx, args.output_md)

        # 可选：保存 JSON
        if args.output_json:
            print(f"\n保存 JSON 数据: {args.output_json}")
            with open(args.output_json, 'w', encoding='utf-8') as f:
                json.dump({'rows': rows_data}, f, ensure_ascii=False, indent=2)
            print(f"✓ JSON 已保存: {args.output_json} ({len(rows_data)} rows)")

        print("\n" + "=" * 80)
        print(f"✓ 完成！共提取 {len(rows_data)} 行数据")
        print("=" * 80)
        print("\n下一步:")
        print(f"  1. 查看 Markdown: cat {args.output_md}")
        print(f"  2. 使用 generate_translation_mapping.py 生成对照表")

        return 0

    except Exception as e:
        print(f"\n✗ 错误: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())
