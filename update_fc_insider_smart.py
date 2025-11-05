#!/usr/bin/env python3
"""
方案 2: 智能过滤 - 读取所有内容后智能处理

假设：
- 读取单元格的所有文本内容
- 使用智能过滤算法移除占位符
- 支持多种占位符模式
- 支持多行文本处理

使用方法：
python3 update_fc_insider_smart.py \
  --input "input.docx" \
  --translations "translations.json" \
  --output "output.docx" \
  --author "Translator Name" \
  --verbose
"""

import argparse
import json
import sys
from datetime import datetime
from typing import Dict, List, Tuple, Optional
from pathlib import Path
import re

try:
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
except ImportError:
    print("错误：需要安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


def smart_filter_placeholders(text: str, verbose: bool = False) -> str:
    """
    智能过滤占位符

    处理多种占位符模式：
    - "<0/>"在第 <1/> 頁
    - "<2/>"
    - 带引号的占位符："<0/>"
    - 独立的占位符：<1/>

    保留真实内容
    """
    original = text

    # 1. 移除引号包裹的占位符：  "<0/>"
    text = re.sub(r'"<\d+/>"', '', text)

    # 2. 移除独立的占位符： <1/>
    text = re.sub(r'<\d+/>', '', text)

    # 3. 清理多余的空白
    # 移除行首尾空白
    lines = [line.strip() for line in text.split('\n')]
    # 移除空行
    lines = [line for line in lines if line]
    # 重新组合
    text = '\n'.join(lines)

    # 4. 清理多余的空格
    text = re.sub(r'\s+', ' ', text).strip()

    if verbose and original != text:
        print(f"    过滤前: '{original[:80]}...'")
        print(f"    过滤后: '{text[:80]}...'")

    return text


def get_cell_text_smart(cell, verbose: bool = False) -> str:
    """
    智能读取单元格文本

    方法：
    1. 读取所有段落的所有文本
    2. 应用智能过滤
    3. 返回清理后的文本
    """
    # 收集所有文本
    all_text_parts = []

    for paragraph in cell.paragraphs:
        para_text = paragraph.text.strip()
        if para_text:
            all_text_parts.append(para_text)

    # 组合所有段落
    full_text = '\n'.join(all_text_parts)

    # 应用智能过滤
    filtered_text = smart_filter_placeholders(full_text, verbose)

    return filtered_text


def get_cell_text_by_style_smart(cell, target_style: Optional[str] = None, verbose: bool = False) -> str:
    """
    按样式智能读取单元格文本

    Args:
        target_style: None = 读取所有样式
                     "Tag" = 只读取 Tag 样式
                     "NotTag" = 只读取非 Tag 样式
    """
    all_text_parts = []

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run_element = run._element
            rpr = run_element.find(qn('w:rPr'))

            has_tag_style = False
            if rpr is not None:
                r_style = rpr.find(qn('w:rStyle'))
                if r_style is not None:
                    style_val = r_style.get(qn('w:val'))
                    if style_val == 'Tag':
                        has_tag_style = True

            # 根据 target_style 决定是否包含
            include = False
            if target_style is None:
                include = True
            elif target_style == "Tag" and has_tag_style:
                include = True
            elif target_style == "NotTag" and not has_tag_style:
                include = True

            if include and run.text:
                all_text_parts.append(run.text)

    # 组合文本
    full_text = ''.join(all_text_parts)

    # 应用智能过滤
    filtered_text = smart_filter_placeholders(full_text, verbose)

    return filtered_text


def analyze_cell_structure(cell) -> Dict:
    """
    分析单元格结构

    返回详细的结构信息，帮助调试
    """
    structure = {
        'paragraphs': [],
        'has_tag_style': False,
        'has_non_tag_style': False,
        'total_runs': 0
    }

    for para_idx, paragraph in enumerate(cell.paragraphs):
        para_info = {
            'index': para_idx,
            'text': paragraph.text,
            'runs': []
        }

        for run_idx, run in enumerate(paragraph.runs):
            run_element = run._element
            rpr = run_element.find(qn('w:rPr'))

            style_val = None
            if rpr is not None:
                r_style = rpr.find(qn('w:rStyle'))
                if r_style is not None:
                    style_val = r_style.get(qn('w:val'))

            run_info = {
                'index': run_idx,
                'text': run.text,
                'style': style_val,
                'is_tag_style': style_val == 'Tag'
            }

            para_info['runs'].append(run_info)
            structure['total_runs'] += 1

            if style_val == 'Tag':
                structure['has_tag_style'] = True
            else:
                structure['has_non_tag_style'] = True

        structure['paragraphs'].append(para_info)

    return structure


def has_track_changes_enabled(doc) -> bool:
    """检查文档是否已启用追踪修订"""
    try:
        settings = doc.settings.element
        track_revisions = settings.find(qn('w:trackRevisions'))
        return track_revisions is not None
    except:
        return False


def enable_track_changes(doc):
    """启用文档层级的追踪修订"""
    try:
        settings = doc.settings.element
        track_revisions = settings.find(qn('w:trackRevisions'))
        if track_revisions is None:
            track_revisions = parse_xml('<w:trackRevisions {} />'.format(
                'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ))
            settings.append(track_revisions)
    except Exception as e:
        print(f"⚠ 警告：无法启用文档层级追踪修订: {e}")


def replace_cell_with_track_changes_smart(
    cell,
    old_text: str,
    new_text: str,
    author: str,
    date_str: str,
    revision_id: int,
    reading_strategy: str = 'all',
    verbose: bool = False
) -> bool:
    """
    使用追踪修订替换单元格内容

    Args:
        reading_strategy:
            'all' - 读取所有文本后过滤
            'tag_only' - 只读取 Tag 样式后过滤
            'non_tag_only' - 只读取非 Tag 样式后过滤
    """
    # 根据策略选择读取方法
    if reading_strategy == 'all':
        current_text = get_cell_text_smart(cell, verbose)
    elif reading_strategy == 'tag_only':
        current_text = get_cell_text_by_style_smart(cell, "Tag", verbose)
    elif reading_strategy == 'non_tag_only':
        current_text = get_cell_text_by_style_smart(cell, "NotTag", verbose)
    else:
        current_text = get_cell_text_smart(cell, verbose)

    # 验证
    if current_text != old_text:
        print(f"  ✗ 文本不匹配")
        print(f"    预期: '{old_text[:100]}...'")
        print(f"    实际: '{current_text[:100]}...'")

        if verbose:
            print(f"\n  === 单元格结构分析 ===")
            structure = analyze_cell_structure(cell)
            print(f"  总 runs: {structure['total_runs']}")
            print(f"  有 Tag 样式: {structure['has_tag_style']}")
            print(f"  有非 Tag 样式: {structure['has_non_tag_style']}")
            for para in structure['paragraphs']:
                print(f"\n  段落 {para['index']}: {para['text'][:50]}")
                for run in para['runs']:
                    print(f"    Run {run['index']}: 样式={run['style']}, 文本='{run['text'][:30]}'")

        return False

    # 清空单元格
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)

    # 确保至少有一个段落
    if not cell.paragraphs:
        cell.add_paragraph()

    paragraph = cell.paragraphs[0]

    # 添加删除标记（旧文本）
    del_run = paragraph.add_run(old_text)
    del_run_element = del_run._element

    # 包装为删除
    del_element = parse_xml(f'''
        <w:del w:id="{revision_id}" w:author="{author}" w:date="{date_str}"
               xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        </w:del>
    ''')

    parent = del_run_element.getparent()
    parent.remove(del_run_element)
    del_element.append(del_run_element)
    paragraph._element.append(del_element)

    # 添加插入标记（新文本）
    ins_element = parse_xml(f'''
        <w:ins w:id="{revision_id + 1}" w:author="{author}" w:date="{date_str}"
               xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        </w:ins>
    ''')

    ins_run_xml = f'''
        <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:t xml:space="preserve">{new_text}</w:t>
        </w:r>
    '''

    ins_run_element = parse_xml(ins_run_xml)
    ins_element.append(ins_run_element)
    paragraph._element.append(ins_element)

    return True


def find_table(doc) -> Optional:
    """查找文档中的第一个表格"""
    if not doc.tables:
        return None
    return doc.tables[0]


def update_translations(
    input_path: str,
    translations_path: str,
    output_path: str,
    author: str = "Translator",
    verbose: bool = False,
    reading_strategy: str = 'all'
) -> Tuple[int, int]:
    """
    更新翻译

    Args:
        reading_strategy: 'all', 'tag_only', 'non_tag_only'
    """
    # 加载文档
    print(f"\n📖 加载文档: {input_path}")
    doc = Document(input_path)

    # 启用追踪修订
    if has_track_changes_enabled(doc):
        print("✓ 文档层级追踪修订已存在")
    else:
        enable_track_changes(doc)
        print("✓ 已启用文档层级追踪修订")

    # 加载翻译
    with open(translations_path, 'r', encoding='utf-8') as f:
        translations = json.load(f)

    # 查找表格
    table = find_table(doc)
    if not table:
        raise ValueError("文档中未找到表格")

    # 构建 segment_id -> row 映射
    row_map = {}
    for i, row in enumerate(table.rows[1:], start=1):  # 跳过表头
        if len(row.cells) >= 4:
            segment_id = row.cells[0].text.strip()
            if segment_id:
                row_map[segment_id] = i

    print(f"\n{'='*80}")
    print(f"FC Insider 翻译更新 - 方案 2 (智能过滤)")
    print(f"读取策略: {reading_strategy}")
    print(f"作者: {author}")
    print(f"翻译数量: {len(translations)}")
    print(f"{'='*80}")

    success_count = 0
    fail_count = 0
    date_str = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    revision_id = 1000

    print(f"\n开始处理 {len(translations)} 个翻译...")
    print("="*80)

    for idx, translation in enumerate(translations, 1):
        segment_id = translation.get('segment_id')
        old_text = translation.get('old_translation', '').strip()
        new_text = translation.get('new_translation', '').strip()

        print(f"[{idx}/{len(translations)}] 处理 {segment_id}...", end=" ")

        if not segment_id or segment_id not in row_map:
            print(f"✗ Segment ID 未找到")
            fail_count += 1
            continue

        row_idx = row_map[segment_id]
        target_cell = table.rows[row_idx].cells[3]

        if verbose:
            print()

        success = replace_cell_with_track_changes_smart(
            target_cell,
            old_text,
            new_text,
            author,
            date_str,
            revision_id,
            reading_strategy,
            verbose
        )

        if success:
            print("✓" if not verbose else "  ✓ 成功")
            success_count += 1
            revision_id += 2
        else:
            fail_count += 1

    print("="*80)
    print(f"\n{'✓ 更新完成' if fail_count == 0 else '⚠ 更新完成（有失败项）'}: {success_count}/{len(translations)}")
    if fail_count > 0:
        print(f"✗ 失败: {fail_count}")
    print("="*80)

    # 保存
    print(f"\n💾 保存文档: {output_path}")
    doc.save(output_path)
    print("✓ 完成")

    return success_count, fail_count


def main():
    parser = argparse.ArgumentParser(
        description='方案 2: 智能过滤 - 读取所有内容后智能处理',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  # 策略 1：读取所有文本（默认）
  python3 update_fc_insider_smart.py \\
    --input "input.docx" \\
    --translations "translations.json" \\
    --output "output.docx" \\
    --author "Gemini" \\
    --strategy all \\
    --verbose

  # 策略 2：只读取 Tag 样式文本
  python3 update_fc_insider_smart.py \\
    --input "input.docx" \\
    --translations "translations.json" \\
    --output "output.docx" \\
    --author "Gemini" \\
    --strategy tag_only \\
    --verbose

  # 策略 3：只读取非 Tag 样式文本
  python3 update_fc_insider_smart.py \\
    --input "input.docx" \\
    --translations "translations.json" \\
    --output "output.docx" \\
    --author "Gemini" \\
    --strategy non_tag_only \\
    --verbose
        """
    )

    parser.add_argument('--input', required=True, help='输入 Word 文档路径')
    parser.add_argument('--translations', required=True, help='翻译映射 JSON 文件路径')
    parser.add_argument('--output', required=True, help='输出 Word 文档路径')
    parser.add_argument('--author', default='Translator', help='追踪修订作者名称')
    parser.add_argument('--strategy',
                       choices=['all', 'tag_only', 'non_tag_only'],
                       default='all',
                       help='读取策略')
    parser.add_argument('--verbose', action='store_true', help='显示详细信息（包括结构分析）')

    args = parser.parse_args()

    try:
        success, fail = update_translations(
            args.input,
            args.translations,
            args.output,
            args.author,
            args.verbose,
            args.strategy
        )

        sys.exit(0 if fail == 0 else 1)

    except Exception as e:
        print(f"\n❌ 错误: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
