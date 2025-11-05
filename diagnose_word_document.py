#!/usr/bin/env python3
"""
诊断 Word 文档单元格读取问题
检查为什么所有 Target 单元格都读取为空
"""

import sys
import argparse
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("错误：需要安装 python-docx")
    print("运行: pip install --user python-docx")
    sys.exit(1)


def find_row_by_segment_id(table, segment_id):
    """在表格中查找包含指定 segment_id 的行"""
    for row_idx, row in enumerate(table.rows):
        first_cell_text = row.cells[0].text.strip()
        if first_cell_text == str(segment_id).strip():
            return row, row_idx
    return None, None


def diagnose_document(docx_path: str, segment_ids: list):
    """
    诊断 Word 文档的表格结构

    Args:
        docx_path: Word 文档路径
        segment_ids: 要检查的 segment ID 列表
    """
    print("=" * 80)
    print("Word 文档诊断")
    print("=" * 80)

    # 打开文档
    print(f"\n读取文档: {docx_path}")
    doc = Document(docx_path)

    # 检查表格
    if not doc.tables:
        print("✗ 文档中没有找到表格！")
        return

    print(f"✓ 找到 {len(doc.tables)} 个表格")

    # 分析第一个表格
    table = doc.tables[0]
    print(f"\n表格 1 分析:")
    print(f"  行数: {len(table.rows)}")
    print(f"  列数: {len(table.rows[0].cells) if table.rows else 0}")

    # 检查表头
    if table.rows:
        print(f"\n表头行:")
        header_row = table.rows[0]
        for i, cell in enumerate(header_row.cells):
            print(f"  列 {i+1}: '{cell.text.strip()}'")

    # 检查指定的 segment IDs
    print("\n" + "=" * 80)
    print("检查指定的 Segment IDs")
    print("=" * 80)

    for segment_id in segment_ids[:5]:  # 只检查前 5 个
        print(f"\n检查 Segment ID: {segment_id}")

        row, row_idx = find_row_by_segment_id(table, segment_id)

        if row is None:
            print(f"  ✗ 未找到此 segment ID")
            continue

        print(f"  ✓ 找到，位于第 {row_idx + 1} 行")

        # 检查所有列
        print(f"  列数: {len(row.cells)}")

        for i, cell in enumerate(row.cells):
            text = cell.text.strip()
            print(f"  列 {i+1} ({['Segment ID', 'Status', 'Source', 'Target'][i] if i < 4 else f'列{i+1}'}): '{text[:100]}'")

            # 特别检查 Target 列（第 4 列，索引 3）
            if i == 3:
                if not text:
                    print(f"    ⚠️  Target 列为空！")
                    print(f"    段落数: {len(cell.paragraphs)}")
                    for p_idx, para in enumerate(cell.paragraphs):
                        print(f"    段落 {p_idx + 1}: '{para.text[:100]}'")
                        print(f"      Runs 数: {len(para.runs)}")
                        for r_idx, run in enumerate(para.runs):
                            print(f"      Run {r_idx + 1}: '{run.text[:50]}'")
                else:
                    print(f"    ✓ Target 列有内容")

    # 随机检查几行数据
    print("\n" + "=" * 80)
    print("随机抽样检查（前 3 行数据）")
    print("=" * 80)

    data_rows = table.rows[1:4] if len(table.rows) > 1 else []

    for row_idx, row in enumerate(data_rows, 1):
        print(f"\n数据行 {row_idx}:")
        cells = row.cells
        print(f"  列数: {len(cells)}")

        for i, cell in enumerate(cells[:4]):  # 只检查前 4 列
            text = cell.text.strip()
            col_name = ['Segment ID', 'Status', 'Source', 'Target'][i] if i < 4 else f'列{i+1}'
            print(f"  {col_name}: '{text[:80]}{'...' if len(text) > 80 else ''}'")

            # 检查 Target 列的详细信息
            if i == 3:
                if not text:
                    print(f"    ⚠️  内容为空")
                    print(f"    段落数: {len(cell.paragraphs)}")
                    print(f"    检查 XML 结构...")

                    # 检查是否有隐藏内容
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text:
                                print(f"    发现 Run 文本: '{run.text[:50]}'")

    # 诊断建议
    print("\n" + "=" * 80)
    print("诊断结果")
    print("=" * 80)

    target_empty_count = 0
    for row in table.rows[1:]:  # 跳过表头
        if len(row.cells) >= 4:
            target_text = row.cells[3].text.strip()
            if not target_text:
                target_empty_count += 1

    total_data_rows = len(table.rows) - 1  # 减去表头

    print(f"\n总数据行: {total_data_rows}")
    print(f"Target 列为空的行: {target_empty_count}")
    print(f"Target 列有内容的行: {total_data_rows - target_empty_count}")

    if target_empty_count == total_data_rows:
        print("\n⚠️  所有 Target 列都为空！")
        print("\n可能原因:")
        print("  1. 这是一个新文档，Target 列还没有翻译")
        print("  2. Target 列使用了特殊格式（如文本框、域代码）")
        print("  3. python-docx 无法正确读取此格式")
        print("\n建议:")
        print("  1. 在 Word 中打开文档，确认 Target 列确实有内容")
        print("  2. 如果有内容，尝试复制粘贴为纯文本")
        print("  3. 或使用 MarkItDown 提取的 Markdown 作为参考")
    elif target_empty_count > 0:
        print(f"\n⚠️  有 {target_empty_count} 行的 Target 列为空")
        print("\n建议:")
        print("  1. 检查这些空行是否应该有内容")
        print("  2. 在对照表中设置 old_text 为空字符串")
    else:
        print("\n✓ 所有 Target 列都有内容")
        print("\n问题可能是:")
        print("  1. old_text 与实际内容不完全匹配（空格、标点等）")
        print("  2. 需要使用模糊匹配而不是精确匹配")


def main():
    parser = argparse.ArgumentParser(
        description='诊断 Word 文档单元格读取问题',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  # 检查文档和对照表
  python diagnose_word_document.py input.docx translations.json

  # 只检查文档
  python diagnose_word_document.py input.docx
        '''
    )

    parser.add_argument('docx_file', help='Word 文档路径')
    parser.add_argument('translations_json', nargs='?', help='翻译对照表（可选）')

    args = parser.parse_args()

    # 检查文件
    if not Path(args.docx_file).exists():
        print(f"✗ 错误：文件不存在 - {args.docx_file}")
        return 1

    # 获取要检查的 segment IDs
    segment_ids = []

    if args.translations_json and Path(args.translations_json).exists():
        import json
        with open(args.translations_json, 'r', encoding='utf-8') as f:
            data = json.load(f)
            translations = data.get('translations', data)
            segment_ids = [t['segment_id'] for t in translations]
            print(f"从对照表读取 {len(segment_ids)} 个 Segment IDs")
    else:
        print("未提供对照表，将检查文档的一般结构")

    try:
        diagnose_document(args.docx_file, segment_ids)
        return 0

    except Exception as e:
        print(f"\n✗ 错误: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())
