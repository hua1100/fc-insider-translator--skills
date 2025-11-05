#!/usr/bin/env python3
"""
方案 1: 反转逻辑 - 只读取 Tag 样式的文本

假设：
- 实际翻译内容带有 Tag 样式
- 占位符（如 <0/>, <1/>）没有 Tag 样式或者需要被过滤掉

使用方法：
python3 update_fc_insider_reverse.py \
  --input "input.docx" \
  --translations "translations.json" \
  --output "output.docx" \
  --author "Translator Name" \
  --verbose
"""

import argparse
import json
import sys
from datetime import datetime
from typing import Dict, List, Tuple, Optional
from pathlib import Path
import re

try:
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
except ImportError:
    print("错误：需要安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


def get_cell_text_only_tags(cell) -> str:
    """
    【方案 1 逻辑】只读取 Tag 样式的文本

    假设真正的翻译内容才带有 Tag 样式
    """
    text_parts = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run_element = run._element
            rpr = run_element.find(qn('w:rPr'))

            # 只保留带有 Tag 样式的 runs
            if rpr is not None:
                r_style = rpr.find(qn('w:rStyle'))
                if r_style is not None:
                    style_val = r_style.get(qn('w:val'))
                    if style_val == 'Tag':
                        if run.text:
                            text_parts.append(run.text)

    full_text = ''.join(text_parts).strip()

    # 过滤掉占位符模式 <数字/>
    # 移除所有 <0/>, <1/>, <2/> 等占位符
    filtered_text = re.sub(r'<\d+/>', '', full_text)

    return filtered_text.strip()


def get_all_cell_text_with_filtering(cell) -> str:
    """
    【方案 1 备选】读取所有文本，但过滤占位符

    可能所有内容都是 Tag 样式，需要通过正则过滤
    """
    full_text = cell.text.strip()

    # 移除占位符模式
    # 处理 "<0/>"在第 <1/> 頁 这种情况
    # 移除引号包裹的占位符
    filtered_text = re.sub(r'"<\d+/>"', '', full_text)
    # 移除独立的占位符
    filtered_text = re.sub(r'<\d+/>', '', filtered_text)

    return filtered_text.strip()


def has_track_changes_enabled(doc) -> bool:
    """检查文档是否已启用追踪修订"""
    try:
        settings = doc.settings.element
        track_revisions = settings.find(qn('w:trackRevisions'))
        return track_revisions is not None
    except:
        return False


def enable_track_changes(doc):
    """启用文档层级的追踪修订"""
    try:
        settings = doc.settings.element
        track_revisions = settings.find(qn('w:trackRevisions'))
        if track_revisions is None:
            track_revisions = parse_xml('<w:trackRevisions {} />'.format(
                'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ))
            settings.append(track_revisions)
    except Exception as e:
        print(f"⚠ 警告：无法启用文档层级追踪修订: {e}")


def preserve_tag_runs(cell) -> List:
    """
    保存单元格中所有 Tag 样式的占位符 runs
    返回需要保留的 runs 信息
    """
    tag_runs = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run_element = run._element
            rpr = run_element.find(qn('w:rPr'))

            if rpr is not None:
                r_style = rpr.find(qn('w:rStyle'))
                if r_style is not None:
                    style_val = r_style.get(qn('w:val'))
                    if style_val == 'Tag':
                        # 如果是占位符（<0/> 格式），保存
                        if run.text and re.match(r'<\d+/>', run.text.strip()):
                            tag_runs.append({
                                'text': run.text,
                                'style': 'Tag'
                            })

    return tag_runs


def replace_cell_with_track_changes(
    cell,
    old_text: str,
    new_text: str,
    author: str,
    date_str: str,
    revision_id: int,
    reading_method: str = 'only_tags'
) -> bool:
    """
    使用追踪修订替换单元格内容

    Args:
        reading_method: 'only_tags' 或 'all_filtered'
    """
    # 根据方法选择读取函数
    if reading_method == 'only_tags':
        current_text = get_cell_text_only_tags(cell)
    else:
        current_text = get_all_cell_text_with_filtering(cell)

    # 保存占位符 Tag runs
    tag_runs = preserve_tag_runs(cell)

    # 验证
    if current_text != old_text:
        print(f"  ✗ 文本不匹配")
        print(f"    预期: '{old_text[:100]}...'")
        print(f"    实际: '{current_text[:100]}...'")
        return False

    # 清空单元格
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)

    # 确保至少有一个段落
    if not cell.paragraphs:
        cell.add_paragraph()

    paragraph = cell.paragraphs[0]

    # 1. 先添加保留的占位符 Tag runs
    for tag_info in tag_runs:
        tag_run = paragraph.add_run(tag_info['text'])
        tag_run.style = tag_info['style']

    # 2. 添加删除标记（旧文本）
    del_run = paragraph.add_run(old_text)
    del_run_element = del_run._element

    # 如果旧文本有 Tag 样式，也保留
    try:
        rpr = del_run_element.find(qn('w:rPr'))
        if rpr is None:
            rpr = parse_xml('<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            del_run_element.insert(0, rpr)

        r_style = parse_xml('<w:rStyle w:val="Tag" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        rpr.append(r_style)
    except:
        pass

    # 包装为删除
    del_element = parse_xml(f'''
        <w:del w:id="{revision_id}" w:author="{author}" w:date="{date_str}"
               xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        </w:del>
    ''')

    parent = del_run_element.getparent()
    parent.remove(del_run_element)
    del_element.append(del_run_element)
    paragraph._element.append(del_element)

    # 3. 添加插入标记（新文本）
    ins_element = parse_xml(f'''
        <w:ins w:id="{revision_id + 1}" w:author="{author}" w:date="{date_str}"
               xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        </w:ins>
    ''')

    ins_run_xml = f'''
        <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:rPr>
                <w:rStyle w:val="Tag"/>
            </w:rPr>
            <w:t xml:space="preserve">{new_text}</w:t>
        </w:r>
    '''

    ins_run_element = parse_xml(ins_run_xml)
    ins_element.append(ins_run_element)
    paragraph._element.append(ins_element)

    return True


def find_table(doc) -> Optional:
    """查找文档中的第一个表格"""
    if not doc.tables:
        return None
    return doc.tables[0]


def update_translations(
    input_path: str,
    translations_path: str,
    output_path: str,
    author: str = "Translator",
    verbose: bool = False,
    reading_method: str = 'only_tags'
) -> Tuple[int, int]:
    """
    Args:
        reading_method: 'only_tags' - 只读 Tag 样式文本
                       'all_filtered' - 读所有文本后过滤
    """
    # 加载文档
    print(f"\n📖 加载文档: {input_path}")
    doc = Document(input_path)

    # 启用追踪修订
    if has_track_changes_enabled(doc):
        print("✓ 文档层级追踪修订已存在")
    else:
        enable_track_changes(doc)
        print("✓ 已启用文档层级追踪修订")

    # 加载翻译
    with open(translations_path, 'r', encoding='utf-8') as f:
        translations = json.load(f)

    # 查找表格
    table = find_table(doc)
    if not table:
        raise ValueError("文档中未找到表格")

    # 构建 segment_id -> row 映射
    row_map = {}
    for i, row in enumerate(table.rows[1:], start=1):  # 跳过表头
        if len(row.cells) >= 4:
            segment_id = row.cells[0].text.strip()
            if segment_id:
                row_map[segment_id] = i

    print(f"\n{'='*80}")
    print(f"FC Insider 翻译更新 - 方案 1 (反转逻辑)")
    print(f"读取方法: {reading_method}")
    print(f"作者: {author}")
    print(f"翻译数量: {len(translations)}")
    print(f"{'='*80}")

    success_count = 0
    fail_count = 0
    date_str = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    revision_id = 1000

    print(f"\n开始处理 {len(translations)} 个翻译...")
    print("="*80)

    for idx, translation in enumerate(translations, 1):
        segment_id = translation.get('segment_id')
        old_text = translation.get('old_translation', '').strip()
        new_text = translation.get('new_translation', '').strip()

        print(f"[{idx}/{len(translations)}] 处理 {segment_id}...", end=" ")

        if not segment_id or segment_id not in row_map:
            print(f"✗ Segment ID 未找到")
            fail_count += 1
            continue

        row_idx = row_map[segment_id]
        target_cell = table.rows[row_idx].cells[3]

        if verbose:
            print()
            if reading_method == 'only_tags':
                current = get_cell_text_only_tags(target_cell)
            else:
                current = get_all_cell_text_with_filtering(target_cell)
            print(f"  预期 old_text: '{old_text[:50]}...'")
            print(f"  实际文本: '{current[:50]}...'")

        success = replace_cell_with_track_changes(
            target_cell,
            old_text,
            new_text,
            author,
            date_str,
            revision_id,
            reading_method
        )

        if success:
            print("✓" if not verbose else "  ✓ 成功")
            success_count += 1
            revision_id += 2
        else:
            fail_count += 1

    print("="*80)
    print(f"\n{'✓ 更新完成' if fail_count == 0 else '⚠ 更新完成（有失败项）'}: {success_count}/{len(translations)}")
    if fail_count > 0:
        print(f"✗ 失败: {fail_count}")
    print("="*80)

    # 保存
    print(f"\n💾 保存文档: {output_path}")
    doc.save(output_path)
    print("✓ 完成")

    return success_count, fail_count


def main():
    parser = argparse.ArgumentParser(
        description='方案 1: 反转逻辑 - 只读取 Tag 样式的文本',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  # 方法 1：只读取 Tag 样式文本（过滤占位符）
  python3 update_fc_insider_reverse.py \\
    --input "input.docx" \\
    --translations "translations.json" \\
    --output "output.docx" \\
    --author "Gemini" \\
    --method only_tags \\
    --verbose

  # 方法 2：读取所有文本后过滤占位符
  python3 update_fc_insider_reverse.py \\
    --input "input.docx" \\
    --translations "translations.json" \\
    --output "output.docx" \\
    --author "Gemini" \\
    --method all_filtered \\
    --verbose
        """
    )

    parser.add_argument('--input', required=True, help='输入 Word 文档路径')
    parser.add_argument('--translations', required=True, help='翻译映射 JSON 文件路径')
    parser.add_argument('--output', required=True, help='输出 Word 文档路径')
    parser.add_argument('--author', default='Translator', help='追踪修订作者名称')
    parser.add_argument('--method',
                       choices=['only_tags', 'all_filtered'],
                       default='only_tags',
                       help='读取方法: only_tags(只读Tag样式) 或 all_filtered(读所有后过滤)')
    parser.add_argument('--verbose', action='store_true', help='显示详细信息')

    args = parser.parse_args()

    try:
        success, fail = update_translations(
            args.input,
            args.translations,
            args.output,
            args.author,
            args.verbose,
            args.method
        )

        sys.exit(0 if fail == 0 else 1)

    except Exception as e:
        print(f"\n❌ 错误: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
