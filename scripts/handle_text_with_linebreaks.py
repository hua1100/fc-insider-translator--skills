#!/usr/bin/env python3
"""
處理包含換行符的文本追踪修訂

此腳本用於處理 Word 文檔中包含內嵌換行符（<w:br/>）的翻譯更新。
當標準工作流程無法處理換行符時，使用此腳本。

使用方法：
python3 handle_text_with_linebreaks.py \
  --input "input.docx" \
  --translations "translations.json" \
  --output "output.docx" \
  --author "Claire.lee@amway.com"
"""

import argparse
import json
import sys
from datetime import datetime
from typing import Dict, List
from pathlib import Path
from lxml import etree

try:
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
except ImportError:
    print("錯誤：需要安裝 python-docx 和 lxml")
    print("運行: pip install python-docx lxml")
    sys.exit(1)


def xml_escape(text: str) -> str:
    """轉義 XML 特殊字符"""
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&apos;'))


def create_run_with_linebreaks(text: str) -> str:
    """
    創建包含換行符的 run XML

    將文本中的 \n 轉換為 <w:br/> 標籤
    """
    parts = text.split('\n')
    run_xml = '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'

    for i, part in enumerate(parts):
        if i > 0:
            # 添加換行符標籤
            run_xml += '<w:br/>'
        if part:  # 只有當部分不為空時才添加文本
            run_xml += f'<w:t xml:space="preserve">{xml_escape(part)}</w:t>'

    run_xml += '</w:r>'
    return run_xml


def clear_cell_tracked_changes(cell):
    """清除單元格中的所有追踪修訂標記"""
    for paragraph in cell.paragraphs:
        para_element = paragraph._element

        # 處理 <w:del> - 完全移除
        del_elements = para_element.findall(qn('w:del'))
        for del_elem in del_elements:
            para_element.remove(del_elem)

        # 處理 <w:ins> - 移除包裝，保留內容
        ins_elements = para_element.findall(qn('w:ins'))
        for ins_elem in ins_elements:
            # 將 <w:ins> 中的 <w:r> 移到段落級別
            runs = ins_elem.findall(qn('w:r'))
            insert_position = para_element.index(ins_elem)

            for run in runs:
                para_element.insert(insert_position, run)
                insert_position += 1

            # 移除 <w:ins> 包裝
            para_element.remove(ins_elem)


def apply_tracked_change_with_linebreaks(
    cell,
    old_text: str,
    new_text: str,
    author: str,
    date_str: str,
    revision_id: int,
    verbose: bool = False
) -> bool:
    """
    應用包含換行符的追踪修訂

    Args:
        cell: Word 表格單元格
        old_text: 舊文本（可能包含 \n）
        new_text: 新文本（可能包含 \n）
        author: 作者名稱
        date_str: 日期字符串
        revision_id: 修訂ID
        verbose: 詳細模式
    """
    # 清除現有追踪修訂
    clear_cell_tracked_changes(cell)

    # 清空單元格
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)

    # 確保至少有一個段落
    if not cell.paragraphs:
        cell.add_paragraph()

    paragraph = cell.paragraphs[0]

    # 創建刪除標記（舊文本）
    del_element = parse_xml(f'''
        <w:del w:id="{revision_id}" w:author="{xml_escape(author)}" w:date="{date_str}"
               xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        </w:del>
    ''')

    # 創建包含換行符的舊文本 run
    old_run_xml = create_run_with_linebreaks(old_text)
    old_run_element = parse_xml(old_run_xml)

    # 修改 run 以使用 delText
    for t_elem in old_run_element.findall('.//' + qn('w:t')):
        t_elem.tag = qn('w:delText')

    del_element.append(old_run_element)
    paragraph._element.append(del_element)

    # 創建插入標記（新文本）
    ins_element = parse_xml(f'''
        <w:ins w:id="{revision_id + 1}" w:author="{xml_escape(author)}" w:date="{date_str}"
               xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        </w:ins>
    ''')

    # 創建包含換行符的新文本 run
    new_run_xml = create_run_with_linebreaks(new_text)
    new_run_element = parse_xml(new_run_xml)

    ins_element.append(new_run_element)
    paragraph._element.append(ins_element)

    if verbose:
        print(f"  ✓ 已應用追踪修訂（包含換行符）")
        if '\n' in old_text:
            print(f"    舊文本包含 {old_text.count(chr(10))} 個換行符")
        if '\n' in new_text:
            print(f"    新文本包含 {new_text.count(chr(10))} 個換行符")

    return True


def find_cell_by_segment_id(table, segment_id: str):
    """根據 Segment ID 查找單元格"""
    for i, row in enumerate(table.rows):
        if len(row.cells) < 1:
            continue

        cell_text = row.cells[0].text.strip()
        if cell_text == segment_id:
            if len(row.cells) >= 3:
                return row.cells[2], i  # 返回 Target 列

    return None, None


def process_translations(
    input_path: str,
    translations_path: str,
    output_path: str,
    author: str,
    verbose: bool = False
):
    """處理翻譯更新"""
    print(f"📖 加載文檔: {input_path}")
    doc = Document(input_path)

    if not doc.tables:
        print("❌ 錯誤：文檔中沒有表格")
        return False

    table = doc.tables[0]

    # 加載翻譯映射
    with open(translations_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
        translations = data.get('translations', data) if isinstance(data, dict) else data

    print(f"✓ 加載 {len(translations)} 個翻譯")
    print()

    # 生成日期和修訂 ID
    date_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
    base_revision_id = 1000

    success_count = 0
    total_count = len(translations)

    print(f"開始處理 {total_count} 個翻譯...")
    print("=" * 80)

    for idx, translation in enumerate(translations, 1):
        segment_id = translation.get('segment_id', '')
        old_text = translation.get('old_text', translation.get('old_translation', ''))
        new_text = translation.get('new_text', translation.get('new_translation', ''))

        if verbose:
            print(f"[{idx}/{total_count}] 處理 {segment_id}...")

        cell, row_idx = find_cell_by_segment_id(table, segment_id)

        if cell is None:
            print(f"  ✗ 找不到 Segment ID: {segment_id}")
            continue

        # 應用追踪修訂（處理換行符）
        revision_id = base_revision_id + (idx * 2)
        success = apply_tracked_change_with_linebreaks(
            cell,
            old_text,
            new_text,
            author,
            date_str,
            revision_id,
            verbose
        )

        if success:
            success_count += 1
            if not verbose:
                print(f"[{idx}/{total_count}] ✓ {segment_id}")

    print("=" * 80)
    print(f"✓ 更新完成: {success_count}/{total_count}")
    print()

    # 保存文檔
    print(f"💾 保存文檔: {output_path}")
    doc.save(output_path)
    print("✓ 完成")

    return True


def main():
    parser = argparse.ArgumentParser(
        description='處理包含換行符的 Word 文檔翻譯更新',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  # 基本用法
  python3 handle_text_with_linebreaks.py \\
    --input "input.docx" \\
    --translations "translations.json" \\
    --output "output.docx"

  # 詳細模式
  python3 handle_text_with_linebreaks.py \\
    --input "input.docx" \\
    --translations "translations.json" \\
    --output "output.docx" \\
    --author "translator@company.com" \\
    --verbose
'''
    )

    parser.add_argument('--input', required=True, help='輸入 Word 文檔路徑')
    parser.add_argument('--translations', required=True, help='翻譯映射 JSON 文件路徑')
    parser.add_argument('--output', required=True, help='輸出 Word 文檔路徑')
    parser.add_argument('--author', default='Claire.lee@amway.com',
                       help='追踪修訂作者名稱（默認：Claire.lee@amway.com）')
    parser.add_argument('--verbose', action='store_true', help='詳細輸出')

    args = parser.parse_args()

    # 驗證文件存在
    if not Path(args.input).exists():
        print(f"❌ 錯誤：輸入文件不存在: {args.input}")
        sys.exit(1)

    if not Path(args.translations).exists():
        print(f"❌ 錯誤：翻譯文件不存在: {args.translations}")
        sys.exit(1)

    print("=" * 80)
    print("處理包含換行符的翻譯更新")
    print("=" * 80)
    print()

    success = process_translations(
        args.input,
        args.translations,
        args.output,
        args.author,
        args.verbose
    )

    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
