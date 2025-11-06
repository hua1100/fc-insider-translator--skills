#!/usr/bin/env python3
"""
方案 3: 深度 XML 分析工具

功能：
1. 完整分析 Word 文档的 XML 结构
2. 显示每个 run 的详细属性（样式、字体、颜色等）
3. 导出原始 XML 以供检查
4. 提供基于实际结构的解决方案建议

使用方法：
python3 analyze_word_structure_deep.py \
  --input "input.docx" \
  --sample-segment "11d76b912e-c3c9-456c-a895-7f4778e6a43f" \
  --export-xml
"""

import argparse
import sys
from pathlib import Path
from typing import Optional, List, Dict
import json

try:
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree
except ImportError:
    print("错误：需要安装 python-docx 和 lxml")
    print("运行: pip install python-docx lxml")
    sys.exit(1)


def get_run_properties(run) -> Dict:
    """
    获取 run 的所有属性

    返回：
    - style: 样式名称
    - bold: 是否粗体
    - italic: 是否斜体
    - underline: 下划线类型
    - color: 文本颜色
    - font_name: 字体名称
    - font_size: 字体大小
    - all_xml_properties: 原始 XML 属性
    """
    properties = {
        'style': None,
        'bold': False,
        'italic': False,
        'underline': None,
        'color': None,
        'font_name': None,
        'font_size': None,
        'all_xml_properties': []
    }

    run_element = run._element
    rpr = run_element.find(qn('w:rPr'))

    if rpr is not None:
        # 样式
        r_style = rpr.find(qn('w:rStyle'))
        if r_style is not None:
            properties['style'] = r_style.get(qn('w:val'))

        # 粗体
        bold = rpr.find(qn('w:b'))
        if bold is not None:
            properties['bold'] = True

        # 斜体
        italic = rpr.find(qn('w:i'))
        if italic is not None:
            properties['italic'] = True

        # 下划线
        underline = rpr.find(qn('w:u'))
        if underline is not None:
            properties['underline'] = underline.get(qn('w:val'))

        # 颜色
        color = rpr.find(qn('w:color'))
        if color is not None:
            properties['color'] = color.get(qn('w:val'))

        # 字体
        r_fonts = rpr.find(qn('w:rFonts'))
        if r_fonts is not None:
            properties['font_name'] = r_fonts.get(qn('w:ascii'))

        # 字体大小
        sz = rpr.find(qn('w:sz'))
        if sz is not None:
            properties['font_size'] = sz.get(qn('w:val'))

        # 收集所有 XML 属性
        for child in rpr:
            tag_name = child.tag.split('}')[-1]  # 去掉命名空间
            attrs = dict(child.attrib)
            properties['all_xml_properties'].append({
                'tag': tag_name,
                'attributes': attrs
            })

    return properties


def analyze_cell_deep(cell, cell_name: str = "单元格") -> Dict:
    """
    深度分析单元格结构

    返回完整的结构信息
    """
    analysis = {
        'cell_name': cell_name,
        'total_paragraphs': len(cell.paragraphs),
        'total_runs': 0,
        'paragraphs': [],
        'summary': {
            'has_tag_style': False,
            'has_non_tag_style': False,
            'tag_style_count': 0,
            'non_tag_style_count': 0,
            'total_text_length': 0,
            'styles_found': set()
        }
    }

    for para_idx, paragraph in enumerate(cell.paragraphs):
        para_info = {
            'index': para_idx,
            'text': paragraph.text,
            'text_length': len(paragraph.text),
            'runs_count': len(paragraph.runs),
            'runs': []
        }

        for run_idx, run in enumerate(paragraph.runs):
            props = get_run_properties(run)

            run_info = {
                'index': run_idx,
                'text': run.text,
                'text_length': len(run.text) if run.text else 0,
                'properties': props
            }

            # 统计
            analysis['total_runs'] += 1
            if run.text:
                analysis['summary']['total_text_length'] += len(run.text)

            if props['style']:
                analysis['summary']['styles_found'].add(props['style'])

            if props['style'] == 'Tag':
                analysis['summary']['has_tag_style'] = True
                analysis['summary']['tag_style_count'] += 1
            else:
                analysis['summary']['has_non_tag_style'] = True
                analysis['summary']['non_tag_style_count'] += 1

            para_info['runs'].append(run_info)

        analysis['paragraphs'].append(para_info)

    # 转换 set 为 list 以便 JSON 序列化
    analysis['summary']['styles_found'] = list(analysis['summary']['styles_found'])

    return analysis


def export_cell_xml(cell, output_path: str):
    """
    导出单元格的原始 XML

    帮助用户查看真实的 XML 结构
    """
    cell_element = cell._element
    xml_string = etree.tostring(
        cell_element,
        pretty_print=True,
        encoding='unicode'
    )

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(xml_string)

    print(f"✓ XML 已导出到: {output_path}")


def print_analysis_report(analysis: Dict, verbose: bool = False):
    """
    打印分析报告
    """
    print(f"\n{'='*80}")
    print(f"单元格分析报告: {analysis['cell_name']}")
    print(f"{'='*80}")

    # 摘要
    summary = analysis['summary']
    print(f"\n📊 摘要:")
    print(f"  总段落数: {analysis['total_paragraphs']}")
    print(f"  总 runs 数: {analysis['total_runs']}")
    print(f"  总文本长度: {summary['total_text_length']} 字符")
    print(f"  Tag 样式 runs: {summary['tag_style_count']}")
    print(f"  非 Tag 样式 runs: {summary['non_tag_style_count']}")
    print(f"  发现的样式: {', '.join(summary['styles_found']) if summary['styles_found'] else '无'}")

    # 详细段落信息
    print(f"\n📝 段落详情:")
    for para in analysis['paragraphs']:
        print(f"\n  段落 {para['index']}:")
        print(f"    文本: '{para['text'][:100]}{'...' if len(para['text']) > 100 else ''}'")
        print(f"    长度: {para['text_length']} 字符")
        print(f"    Runs: {para['runs_count']}")

        if verbose:
            for run in para['runs']:
                print(f"\n      Run {run['index']}:")
                print(f"        文本: '{run['text'][:50] if run['text'] else '(空)'}{'...' if run['text'] and len(run['text']) > 50 else ''}'")
                print(f"        长度: {run['text_length']} 字符")

                props = run['properties']
                print(f"        样式: {props['style'] or '(无)'}")
                if props['bold']:
                    print(f"        粗体: 是")
                if props['italic']:
                    print(f"        斜体: 是")
                if props['underline']:
                    print(f"        下划线: {props['underline']}")
                if props['color']:
                    print(f"        颜色: {props['color']}")
                if props['font_name']:
                    print(f"        字体: {props['font_name']}")
                if props['font_size']:
                    print(f"        大小: {props['font_size']}")

                if props['all_xml_properties']:
                    print(f"        XML 属性:")
                    for xml_prop in props['all_xml_properties']:
                        print(f"          <{xml_prop['tag']}> {xml_prop['attributes']}")


def generate_solution_recommendation(analysis: Dict) -> str:
    """
    基于分析结果生成解决方案建议
    """
    summary = analysis['summary']

    recommendations = []
    recommendations.append("\n" + "="*80)
    recommendations.append("💡 解决方案建议")
    recommendations.append("="*80)

    # 分析文本分布
    if summary['tag_style_count'] == 0:
        recommendations.append("\n⚠️  发现：没有任何 Tag 样式的 runs")
        recommendations.append("   建议：使用 cell.text 直接读取，然后用正则过滤占位符")
        recommendations.append("   脚本：update_fc_insider_smart.py --strategy all")

    elif summary['non_tag_style_count'] == 0:
        recommendations.append("\n✓ 发现：所有 runs 都是 Tag 样式")
        recommendations.append("  建议：读取所有 Tag 样式文本，然后用正则过滤占位符")
        recommendations.append("  脚本：update_fc_insider_reverse.py --method only_tags")
        recommendations.append("  或者：update_fc_insider_smart.py --strategy tag_only")

    else:
        recommendations.append(f"\n✓ 发现：混合样式")
        recommendations.append(f"  - Tag 样式 runs: {summary['tag_style_count']}")
        recommendations.append(f"  - 非 Tag 样式 runs: {summary['non_tag_style_count']}")

        # 检查哪种样式包含更多文本
        tag_text_length = 0
        non_tag_text_length = 0

        for para in analysis['paragraphs']:
            for run in para['runs']:
                if run['properties']['style'] == 'Tag':
                    tag_text_length += run['text_length']
                else:
                    non_tag_text_length += run['text_length']

        recommendations.append(f"  - Tag 样式文本长度: {tag_text_length}")
        recommendations.append(f"  - 非 Tag 样式文本长度: {non_tag_text_length}")

        if tag_text_length > non_tag_text_length:
            recommendations.append("\n  💡 建议：主要内容在 Tag 样式中")
            recommendations.append("     脚本：update_fc_insider_smart.py --strategy tag_only --verbose")
        else:
            recommendations.append("\n  💡 建议：主要内容在非 Tag 样式中")
            recommendations.append("     脚本：update_fc_insider_smart.py --strategy non_tag_only --verbose")

    # 检查是否需要查看 XML
    if summary['total_runs'] > 10:
        recommendations.append("\n⚠️  复杂结构：runs 数量较多")
        recommendations.append("   建议：使用 --export-xml 导出 XML 进行详细检查")

    recommendations.append("\n" + "="*80)

    return '\n'.join(recommendations)


def find_table(doc) -> Optional:
    """查找文档中的第一个表格"""
    if not doc.tables:
        return None
    return doc.tables[0]


def main():
    parser = argparse.ArgumentParser(
        description='方案 3: 深度 XML 分析工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  # 基本分析
  python3 analyze_word_structure_deep.py \\
    --input "input.docx" \\
    --sample-segment "11d76b912e-c3c9-456c-a895-7f4778e6a43f"

  # 详细分析（显示所有 run 属性）
  python3 analyze_word_structure_deep.py \\
    --input "input.docx" \\
    --sample-segment "11d76b912e-c3c9-456c-a895-7f4778e6a43f" \\
    --verbose

  # 导出 XML
  python3 analyze_word_structure_deep.py \\
    --input "input.docx" \\
    --sample-segment "11d76b912e-c3c9-456c-a895-7f4778e6a43f" \\
    --export-xml

  # 导出 JSON
  python3 analyze_word_structure_deep.py \\
    --input "input.docx" \\
    --sample-segment "11d76b912e-c3c9-456c-a895-7f4778e6a43f" \\
    --export-json analysis.json
        """
    )

    parser.add_argument('--input', required=True, help='输入 Word 文档路径')
    parser.add_argument('--sample-segment', help='样本 Segment ID（分析此行）')
    parser.add_argument('--verbose', action='store_true', help='显示详细的 run 属性')
    parser.add_argument('--export-xml', action='store_true', help='导出单元格的原始 XML')
    parser.add_argument('--export-json', help='导出分析结果为 JSON 文件')

    args = parser.parse_args()

    try:
        # 加载文档
        print(f"\n📖 加载文档: {args.input}")
        doc = Document(args.input)

        # 查找表格
        table = find_table(doc)
        if not table:
            print("❌ 错误：文档中未找到表格")
            sys.exit(1)

        print(f"✓ 找到表格，共 {len(table.rows)} 行")

        # 如果指定了 sample_segment，分析该行
        if args.sample_segment:
            # 查找行
            target_row = None
            target_row_idx = None

            for i, row in enumerate(table.rows[1:], start=1):  # 跳过表头
                if len(row.cells) >= 4:
                    segment_id = row.cells[0].text.strip()
                    if segment_id == args.sample_segment:
                        target_row = row
                        target_row_idx = i
                        break

            if not target_row:
                print(f"❌ 错误：未找到 Segment ID: {args.sample_segment}")
                sys.exit(1)

            print(f"✓ 找到目标行: 第 {target_row_idx} 行")

            # 分析 Target 列（第 4 列，索引 3）
            target_cell = target_row.cells[3]
            analysis = analyze_cell_deep(target_cell, f"Target 列 (行 {target_row_idx})")

            # 打印报告
            print_analysis_report(analysis, args.verbose)

            # 生成建议
            recommendation = generate_solution_recommendation(analysis)
            print(recommendation)

            # 导出 XML
            if args.export_xml:
                xml_filename = f"cell_row{target_row_idx}_xml.xml"
                export_cell_xml(target_cell, xml_filename)

            # 导出 JSON
            if args.export_json:
                with open(args.export_json, 'w', encoding='utf-8') as f:
                    json.dump(analysis, f, ensure_ascii=False, indent=2)
                print(f"✓ JSON 已导出到: {args.export_json}")

        else:
            # 没有指定 segment，分析所有行的 Target 列
            print("\n分析所有行的 Target 列...")

            all_analyses = []

            for i, row in enumerate(table.rows[1:], start=1):
                if len(row.cells) >= 4:
                    segment_id = row.cells[0].text.strip()
                    target_cell = row.cells[3]

                    analysis = analyze_cell_deep(target_cell, f"行 {i} ({segment_id})")
                    all_analyses.append(analysis)

            # 打印摘要
            print(f"\n{'='*80}")
            print(f"所有行摘要 ({len(all_analyses)} 行)")
            print(f"{'='*80}")

            for analysis in all_analyses:
                summary = analysis['summary']
                print(f"\n{analysis['cell_name']}:")
                print(f"  Runs: {analysis['total_runs']}, Tag样式: {summary['tag_style_count']}, 非Tag: {summary['non_tag_style_count']}")
                print(f"  文本长度: {summary['total_text_length']}")

            # 导出 JSON
            if args.export_json:
                with open(args.export_json, 'w', encoding='utf-8') as f:
                    json.dump(all_analyses, f, ensure_ascii=False, indent=2)
                print(f"\n✓ JSON 已导出到: {args.export_json}")

    except Exception as e:
        print(f"\n❌ 错误: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
