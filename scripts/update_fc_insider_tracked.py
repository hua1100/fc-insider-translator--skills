#!/usr/bin/env python3
"""
解决方案 4: 处理已包含追踪修订的单元格

发现的问题：
- 单元格已经包含追踪修订（<w:del> 和 <w:ins>）
- python-docx 的 paragraph.runs 无法读取追踪修订中的 runs
- 导致所有读取方法都返回空内容

此脚本提供三种模式：
1. read_deleted - 读取删除的文本（<w:delText>）
2. read_inserted - 读取插入的文本（<w:t> in <w:ins>）
3. clear_and_update - 清除现有追踪修订，重新应用

使用方法：
python3 update_fc_insider_tracked.py \
  --input "input.docx" \
  --translations "translations.json" \
  --output "output.docx" \
  --author "Translator Name" \
  --mode read_deleted \
  --verbose
"""

import argparse
import json
import sys
from datetime import datetime
from typing import Dict, List, Tuple, Optional
from pathlib import Path
import re

try:
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    from lxml import etree
except ImportError:
    print("错误：需要安装 python-docx 和 lxml")
    print("运行: pip install python-docx lxml")
    sys.exit(1)


def get_cell_text_from_tracked_changes(cell, mode: str = 'read_deleted', verbose: bool = False) -> str:
    """
    从追踪修订中读取文本

    Args:
        mode: 'read_deleted' - 读取删除的文本
              'read_inserted' - 读取插入的文本
              'read_both' - 读取两者（先删除，后插入，用换行分隔）
    """
    text_parts = []

    for paragraph in cell.paragraphs:
        para_element = paragraph._element

        if mode == 'read_deleted' or mode == 'read_both':
            # 查找所有 <w:del> 元素
            del_elements = para_element.findall(qn('w:del'))
            for del_elem in del_elements:
                # 在 <w:del> 中查找 <w:delText>
                del_texts = del_elem.findall('.//' + qn('w:delText'))
                for del_text in del_texts:
                    if del_text.text:
                        text_parts.append(del_text.text)

        if mode == 'read_inserted' or mode == 'read_both':
            # 查找所有 <w:ins> 元素
            ins_elements = para_element.findall(qn('w:ins'))
            for ins_elem in ins_elements:
                # 在 <w:ins> 中查找 <w:t>
                t_elements = ins_elem.findall('.//' + qn('w:t'))
                for t_elem in t_elements:
                    if t_elem.text:
                        text_parts.append(t_elem.text)

    full_text = ''.join(text_parts).strip()

    if verbose:
        print(f"    模式 {mode} 读取到: '{full_text[:80]}...'")

    return full_text


def get_cell_text_normal_or_tracked(cell, mode: str = 'auto', verbose: bool = False) -> Tuple[str, str]:
    """
    智能读取单元格文本

    返回: (text, source)
        text: 读取到的文本
        source: 'normal' | 'deleted' | 'inserted' | 'empty'
    """
    # 先尝试普通读取
    normal_text = cell.text.strip()
    if normal_text:
        return (normal_text, 'normal')

    # 尝试从追踪修订读取
    deleted_text = get_cell_text_from_tracked_changes(cell, 'read_deleted', verbose)
    if deleted_text:
        return (deleted_text, 'deleted')

    inserted_text = get_cell_text_from_tracked_changes(cell, 'read_inserted', verbose)
    if inserted_text:
        return (inserted_text, 'inserted')

    return ('', 'empty')


def clear_cell_tracked_changes(cell):
    """
    清除单元格中的所有追踪修订标记

    保留实际内容，移除 <w:del> 和 <w:ins> 包装
    """
    for paragraph in cell.paragraphs:
        para_element = paragraph._element

        # 处理 <w:del> - 完全移除
        del_elements = para_element.findall(qn('w:del'))
        for del_elem in del_elements:
            para_element.remove(del_elem)

        # 处理 <w:ins> - 移除包装，保留内容
        ins_elements = para_element.findall(qn('w:ins'))
        for ins_elem in ins_elements:
            # 将 <w:ins> 中的 <w:r> 移到段落级别
            runs = ins_elem.findall(qn('w:r'))
            insert_position = para_element.index(ins_elem)

            for run in runs:
                ins_elem.remove(run)
                para_element.insert(insert_position, run)
                insert_position += 1

            # 移除空的 <w:ins>
            para_element.remove(ins_elem)


def has_track_changes_enabled(doc) -> bool:
    """检查文档是否已启用追踪修订"""
    try:
        settings = doc.settings.element
        track_revisions = settings.find(qn('w:trackRevisions'))
        return track_revisions is not None
    except:
        return False


def enable_track_changes(doc):
    """启用文档层级的追踪修订"""
    try:
        settings = doc.settings.element
        track_revisions = settings.find(qn('w:trackRevisions'))
        if track_revisions is None:
            track_revisions = parse_xml('<w:trackRevisions {} />'.format(
                'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ))
            settings.append(track_revisions)
    except Exception as e:
        print(f"⚠ 警告：无法启用文档层级追踪修订: {e}")


def replace_cell_with_track_changes_from_tracked(
    cell,
    old_text: str,
    new_text: str,
    author: str,
    date_str: str,
    revision_id: int,
    reading_mode: str = 'read_deleted',
    update_mode: str = 'clear_and_replace',
    verbose: bool = False
) -> bool:
    """
    替换已包含追踪修订的单元格

    Args:
        reading_mode: 'read_deleted' | 'read_inserted' | 'auto'
        update_mode: 'clear_and_replace' - 清除现有追踪修订后替换
                    'keep_and_add' - 保留现有追踪修订，添加新的（不推荐）
    """
    # 读取当前文本
    if reading_mode == 'auto':
        current_text, source = get_cell_text_normal_or_tracked(cell, reading_mode, verbose)
        if verbose:
            print(f"    自动检测到文本来源: {source}")
    else:
        current_text = get_cell_text_from_tracked_changes(cell, reading_mode, verbose)

    # 验证
    if current_text != old_text:
        print(f"  ✗ 文本不匹配")
        print(f"    预期: '{old_text[:100]}...'")
        print(f"    实际: '{current_text[:100]}...'")
        return False

    # 根据更新模式处理
    if update_mode == 'clear_and_replace':
        # 清除所有追踪修订
        clear_cell_tracked_changes(cell)

        # 清空单元格
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run._element.getparent().remove(run._element)

        # 确保至少有一个段落
        if not cell.paragraphs:
            cell.add_paragraph()

        paragraph = cell.paragraphs[0]

        # 添加新的追踪修订
        # 删除标记
        del_run = paragraph.add_run(old_text)
        del_run_element = del_run._element

        del_element = parse_xml(f'''
            <w:del w:id="{revision_id}" w:author="{author}" w:date="{date_str}"
                   xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            </w:del>
        ''')

        parent = del_run_element.getparent()
        parent.remove(del_run_element)
        del_element.append(del_run_element)
        paragraph._element.append(del_element)

        # 插入标记
        ins_element = parse_xml(f'''
            <w:ins w:id="{revision_id + 1}" w:author="{author}" w:date="{date_str}"
                   xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            </w:ins>
        ''')

        ins_run_xml = f'''
            <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:t xml:space="preserve">{new_text}</w:t>
            </w:r>
        '''

        ins_run_element = parse_xml(ins_run_xml)
        ins_element.append(ins_run_element)
        paragraph._element.append(ins_element)

    return True


def find_table(doc) -> Optional:
    """查找文档中的第一个表格"""
    if not doc.tables:
        return None
    return doc.tables[0]


def update_translations(
    input_path: str,
    translations_path: str,
    output_path: str,
    author: str = "Translator",
    verbose: bool = False,
    reading_mode: str = 'auto',
    update_mode: str = 'clear_and_replace'
) -> Tuple[int, int]:
    """
    更新包含追踪修订的翻译

    Args:
        reading_mode: 'auto' | 'read_deleted' | 'read_inserted'
        update_mode: 'clear_and_replace'
    """
    # 加载文档
    print(f"\n📖 加载文档: {input_path}")
    doc = Document(input_path)

    # 启用追踪修订
    if has_track_changes_enabled(doc):
        print("✓ 文档层级追踪修订已存在")
    else:
        enable_track_changes(doc)
        print("✓ 已启用文档层级追踪修订")

    # 加载翻译
    with open(translations_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
        # 提取 translations 数组（如果存在）
        translations = data.get('translations', data) if isinstance(data, dict) else data

    # 查找表格
    table = find_table(doc)
    if not table:
        raise ValueError("文档中未找到表格")

    # 构建 segment_id -> row 映射
    row_map = {}
    for i, row in enumerate(table.rows[1:], start=1):  # 跳过表头
        if len(row.cells) >= 4:
            segment_id = row.cells[0].text.strip()
            if segment_id:
                row_map[segment_id] = i

    print(f"\n{'='*80}")
    print(f"FC Insider 翻译更新 - 方案 4 (处理追踪修订)")
    print(f"读取模式: {reading_mode}")
    print(f"更新模式: {update_mode}")
    print(f"作者: {author}")
    print(f"翻译数量: {len(translations)}")
    print(f"{'='*80}")

    success_count = 0
    fail_count = 0
    date_str = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    revision_id = 1000

    print(f"\n开始处理 {len(translations)} 个翻译...")
    print("="*80)

    for idx, translation in enumerate(translations, 1):
        segment_id = translation.get('segment_id')
        # 支持两种键名：old_text/new_text 或 old_translation/new_translation
        old_text = translation.get('old_text', translation.get('old_translation', '')).strip()
        new_text = translation.get('new_text', translation.get('new_translation', '')).strip()

        print(f"[{idx}/{len(translations)}] 处理 {segment_id}...", end=" ")

        if not segment_id or segment_id not in row_map:
            print(f"✗ Segment ID 未找到")
            fail_count += 1
            continue

        row_idx = row_map[segment_id]
        target_cell = table.rows[row_idx].cells[3]

        if verbose:
            print()

        success = replace_cell_with_track_changes_from_tracked(
            target_cell,
            old_text,
            new_text,
            author,
            date_str,
            revision_id,
            reading_mode,
            update_mode,
            verbose
        )

        if success:
            print("✓" if not verbose else "  ✓ 成功")
            success_count += 1
            revision_id += 2
        else:
            fail_count += 1

    print("="*80)
    print(f"\n{'✓ 更新完成' if fail_count == 0 else '⚠ 更新完成（有失败项）'}: {success_count}/{len(translations)}")
    if fail_count > 0:
        print(f"✗ 失败: {fail_count}")
    print("="*80)

    # 保存
    print(f"\n💾 保存文档: {output_path}")
    doc.save(output_path)
    print("✓ 完成")

    return success_count, fail_count


def main():
    parser = argparse.ArgumentParser(
        description='方案 4: 处理已包含追踪修订的单元格',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  # 自动检测模式（推荐）
  python3 update_fc_insider_tracked.py \\
    --input "input.docx" \\
    --translations "translations.json" \\
    --output "output.docx" \\
    --author "Gemini" \\
    --mode auto \\
    --verbose

  # 从删除的文本中读取（<w:del>）
  python3 update_fc_insider_tracked.py \\
    --input "input.docx" \\
    --translations "translations.json" \\
    --output "output.docx" \\
    --author "Gemini" \\
    --mode read_deleted \\
    --verbose

  # 从插入的文本中读取（<w:ins>）
  python3 update_fc_insider_tracked.py \\
    --input "input.docx" \\
    --translations "translations.json" \\
    --output "output.docx" \\
    --author "Gemini" \\
    --mode read_inserted \\
    --verbose
        """
    )

    parser.add_argument('--input', required=True, help='输入 Word 文档路径')
    parser.add_argument('--translations', required=True, help='翻译映射 JSON 文件路径')
    parser.add_argument('--output', required=True, help='输出 Word 文档路径')
    parser.add_argument('--author', default='Claire.lee@amway.com', help='追踪修订作者名称（默认：Claire.lee@amway.com）')
    parser.add_argument('--mode',
                       choices=['auto', 'read_deleted', 'read_inserted'],
                       default='auto',
                       help='读取模式')
    parser.add_argument('--verbose', action='store_true', help='显示详细信息')

    args = parser.parse_args()

    try:
        success, fail = update_translations(
            args.input,
            args.translations,
            args.output,
            args.author,
            args.verbose,
            args.mode
        )

        sys.exit(0 if fail == 0 else 1)

    except Exception as e:
        print(f"\n❌ 错误: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
