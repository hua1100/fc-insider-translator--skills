#!/usr/bin/env python3
"""
FC Insider 翻译更新脚本（支持模糊匹配）
- 支持宽松的文本匹配（忽略空格、标点差异）
- 更好的错误诊断
- 可选：允许空 old_text
"""

import json
import argparse
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要安装 python-docx")
    print("运行: pip install --user python-docx")
    exit(1)

from tag_protector import protect_tags, restore_tags


def normalize_text(text: str) -> str:
    """
    规范化文本用于比较
    - 移除多余空格
    - 统一标点符号
    """
    if not text:
        return ""

    # 移除多余空格
    text = ' '.join(text.split())

    # 统一标点符号（可选）
    # text = text.replace('，', ',').replace('。', '.')

    return text.strip()


def fuzzy_match(text1: str, text2: str, threshold: float = 0.9) -> bool:
    """
    模糊匹配两个文本

    Args:
        text1: 文本 1
        text2: 文本 2
        threshold: 相似度阈值 (0-1)

    Returns:
        是否匹配
    """
    # 规范化
    norm1 = normalize_text(text1)
    norm2 = normalize_text(text2)

    # 精确匹配
    if norm1 == norm2:
        return True

    # 包含匹配（如果 text1 是 text2 的子串，或反之）
    if norm1 in norm2 or norm2 in norm1:
        return True

    # 简单的相似度计算（Jaccard 相似度）
    words1 = set(norm1.split())
    words2 = set(norm2.split())

    if not words1 or not words2:
        return False

    intersection = words1 & words2
    union = words1 | words2

    similarity = len(intersection) / len(union)

    return similarity >= threshold


def enable_track_changes(doc):
    """启用文档的追踪修订功能"""
    settings = doc.settings
    settings_element = settings.element

    track_revisions = settings_element.find(qn('w:trackRevisions'))
    if track_revisions is None:
        track_revisions = OxmlElement('w:trackRevisions')
        settings_element.append(track_revisions)
        print("✓ 文档层级追踪修订已启用")
    else:
        print("✓ 文档层级追踪修订已存在")


def create_delete_element(author, date, text_content, revision_id, rpr_element=None):
    """创建删除标记元素"""
    del_element = OxmlElement('w:del')
    del_element.set(qn('w:id'), str(revision_id))
    del_element.set(qn('w:author'), author)
    del_element.set(qn('w:date'), date)

    del_run = OxmlElement('w:r')

    if rpr_element is not None:
        del_run.append(rpr_element)

    del_text = OxmlElement('w:delText')
    del_text.text = text_content
    del_text.set(qn('xml:space'), 'preserve')

    del_run.append(del_text)
    del_element.append(del_run)

    return del_element


def create_insert_element(author, date, text_content, revision_id, rpr_element=None):
    """创建插入标记元素"""
    ins_element = OxmlElement('w:ins')
    ins_element.set(qn('w:id'), str(revision_id))
    ins_element.set(qn('w:author'), author)
    ins_element.set(qn('w:date'), date)

    ins_run = OxmlElement('w:r')

    if rpr_element is not None:
        ins_run.append(rpr_element)

    ins_text = OxmlElement('w:t')
    ins_text.text = text_content
    ins_text.set(qn('xml:space'), 'preserve')

    ins_run.append(ins_text)
    ins_element.append(ins_run)

    return ins_element


def find_row_by_segment_id(table, segment_id):
    """在表格中查找包含指定 segment_id 的行"""
    for row_idx, row in enumerate(table.rows):
        first_cell_text = row.cells[0].text.strip()
        if first_cell_text == str(segment_id).strip():
            return row, row_idx
    return None, None


def replace_cell_with_track_changes(cell, old_text, new_text, author, date_str, revision_id, fuzzy: bool = False):
    """
    使用追踪修订替换单元格内容

    Args:
        cell: 要修改的单元格
        old_text: 旧文本（用于验证）
        new_text: 新文本
        author: 作者名称
        date_str: 日期字符串
        revision_id: 修订 ID
        fuzzy: 是否使用模糊匹配

    Returns:
        新的 revision_id
    """
    # 获取当前文本
    current_text = cell.text.strip()

    # 保护标签
    old_protected = protect_tags(old_text)
    new_protected = protect_tags(new_text)
    current_protected = protect_tags(current_text)

    # 验证文本
    allow_empty_old = (old_text is None) or (str(old_text).strip() == "")

    if not allow_empty_old:
        if fuzzy:
            # 模糊匹配
            if not fuzzy_match(old_protected, current_protected):
                raise ValueError(
                    f"文本不匹配（模糊匹配）\n"
                    f"  预期: '{old_text[:50]}...'\n"
                    f"  实际: '{current_text[:50]}...'\n"
                    f"  规范化预期: '{normalize_text(old_text)[:50]}...'\n"
                    f"  规范化实际: '{normalize_text(current_text)[:50]}...'"
                )
        else:
            # 精确匹配
            if old_protected not in current_protected:
                raise ValueError(
                    f"文本不匹配（精确匹配）\n"
                    f"  预期: '{old_text[:50]}...'\n"
                    f"  实际: '{current_text[:50]}...'"
                )

    # 获取第一个段落
    if len(cell.paragraphs) == 0:
        return revision_id

    paragraph = cell.paragraphs[0]
    p_element = paragraph._element

    # 保存格式
    rpr_element = None
    if paragraph.runs:
        first_run = paragraph.runs[0]._element
        rpr_nodes = first_run.findall(qn('w:rPr'))
        if rpr_nodes:
            rpr_element = rpr_nodes[0]

    # 清空段落内容
    for child in list(p_element):
        if child.tag in [qn('w:r'), qn('w:del'), qn('w:ins'),
                         qn('w:hyperlink'), qn('w:bookmarkStart'),
                         qn('w:bookmarkEnd')]:
            p_element.remove(child)

    # 添加删除标记
    if old_protected.strip():
        del_element = create_delete_element(
            author, date_str, old_protected, revision_id, rpr_element
        )
        p_element.append(del_element)
        revision_id += 1

    # 添加插入标记
    if new_protected.strip():
        ins_element = create_insert_element(
            author, date_str, new_protected, revision_id, rpr_element
        )
        p_element.append(ins_element)
        revision_id += 1

    return revision_id


def update_translations(doc, translations, author='Claude', fuzzy=False, verbose=False):
    """更新翻译"""
    successful = 0
    failed = []

    if not doc.tables:
        raise ValueError("文档中没有找到表格")

    table = doc.tables[0]
    date_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')
    revision_id = 0

    print(f"开始处理 {len(translations)} 个翻译...")
    print(f"匹配模式: {'模糊匹配' if fuzzy else '精确匹配'}")
    print("=" * 80)

    for i, trans in enumerate(translations, 1):
        segment_id = trans.get('segment_id', 'Unknown')
        old_text = trans.get('old_text', '')
        new_text = trans.get('new_text', '')

        print(f"[{i}/{len(translations)}] 处理 {segment_id[:40]}...", end=' ')

        if verbose:
            print(f"\n  预期 old_text: '{old_text[:50]}...'")

        try:
            # 查找行
            row, row_idx = find_row_by_segment_id(table, segment_id)
            if row is None:
                raise ValueError(f"找不到 segment_id: {segment_id}")

            # 获取目标单元格
            if len(row.cells) < 4:
                raise ValueError("表格列数不足（需要至少 4 列）")

            target_cell = row.cells[3]

            # 诊断：显示实际内容
            if verbose:
                actual_text = target_cell.text.strip()
                print(f"  实际文本: '{actual_text[:50]}...'")
                print(f"  匹配结果: {fuzzy_match(old_text, actual_text) if fuzzy else (old_text in actual_text)}")

            # 应用追踪修订
            revision_id = replace_cell_with_track_changes(
                target_cell, old_text, new_text, author, date_str, revision_id, fuzzy
            )

            successful += 1
            print("✓")

        except Exception as e:
            error_msg = str(e)
            failed.append((segment_id, error_msg))

            if verbose:
                print(f"✗\n  详细错误: {error_msg}")
            else:
                # 截取错误信息
                short_error = error_msg.split('\n')[0][:60]
                print(f"✗ ({short_error})")

    return successful, failed


def restore_tags_in_document(doc):
    """恢复文档中的标签保护"""
    print("\n恢复标签保护...")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            run.text = restore_tags(run.text)

    print("✓ 标签恢复完成")


def main():
    parser = argparse.ArgumentParser(
        description='FC Insider 翻译更新（支持模糊匹配）',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  # 精确匹配（默认）
  python update_fc_insider_fuzzy.py --input input.docx --translations trans.json --output output.docx

  # 模糊匹配（推荐）
  python update_fc_insider_fuzzy.py --input input.docx --translations trans.json --output output.docx --fuzzy

  # 详细模式（诊断用）
  python update_fc_insider_fuzzy.py --input input.docx --translations trans.json --output output.docx --fuzzy --verbose
        '''
    )

    parser.add_argument('--input', required=True, help='输入 DOCX 文件')
    parser.add_argument('--translations', required=True, help='翻译 JSON 文件路径')
    parser.add_argument('--output', required=True, help='输出 DOCX 文件')
    parser.add_argument('--author', default='Claude', help='作者名称（默认: Claude）')
    parser.add_argument('--fuzzy', action='store_true', help='使用模糊匹配（推荐）')
    parser.add_argument('--verbose', action='store_true', help='显示详细诊断信息')

    args = parser.parse_args()

    # 检查输入文件
    if not Path(args.input).exists():
        print(f"✗ 错误：输入文件不存在 - {args.input}")
        return 1

    if not Path(args.translations).exists():
        print(f"✗ 错误：翻译文件不存在 - {args.translations}")
        return 1

    try:
        # 加载翻译数据
        with open(args.translations, 'r', encoding='utf-8') as f:
            data = json.load(f)
            translations = data.get('translations', data)

        # 打开文档
        print(f"读取文档: {args.input}")
        doc = Document(args.input)

        # 启用追踪修订
        enable_track_changes(doc)

        print(f"FC Insider 翻译更新（模糊匹配版）")
        print(f"作者: {args.author}")
        print(f"翻译数量: {len(translations)}")
        print("=" * 80)

        # 更新翻译
        successful, failed = update_translations(doc, translations, args.author, args.fuzzy, args.verbose)

        # 恢复标签
        restore_tags_in_document(doc)

        print("=" * 80)
        print(f"✓ 更新完成: {successful}/{len(translations)}")

        if failed:
            print(f"✗ 失败: {len(failed)} 个")
            print("\n失敗詳情（前 10 個）:")
            for seg_id, error in failed[:10]:
                print(f"  • {seg_id[:40]}: {error[:200]}")

        print("=" * 80)

        # 保存文档
        print(f"\n保存文档: {args.output}")
        doc.save(args.output)
        print("✓ 文档已保存")

        return 0 if successful > 0 else 1

    except Exception as e:
        print(f"\n✗ 致命错误: {e}")
        import traceback
        traceback.print_exc()
        return 2


if __name__ == '__main__':
    import sys
    sys.exit(main())
