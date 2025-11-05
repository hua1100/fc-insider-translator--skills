#!/usr/bin/env python3
"""
FC Insider 翻译更新脚本（简化版 - 适用于 Claude Skills）
- 直接使用 python-docx，无需 unpack/pack
- 使用追踪修订标记变更
- 处理 FC Insider 四栏表格结构
"""

import json
import argparse
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
except ImportError:
    print("错误：需要安装 python-docx")
    print("运行: pip install --user python-docx")
    exit(1)

from tag_protector import protect_tags, restore_tags


def enable_track_changes(doc):
    """启用文档的追踪修订功能"""
    settings = doc.settings
    settings_element = settings.element

    # 检查是否已有 trackRevisions 元素
    track_revisions = settings_element.find(qn('w:trackRevisions'))
    if track_revisions is None:
        track_revisions = OxmlElement('w:trackRevisions')
        settings_element.append(track_revisions)
        print("✓ 文档层级追踪修订已启用")
    else:
        print("✓ 文档层级追踪修订已存在")


def create_delete_element(author, date, text_content, revision_id, rpr_element=None):
    """创建删除标记元素"""
    del_element = OxmlElement('w:del')
    del_element.set(qn('w:id'), str(revision_id))
    del_element.set(qn('w:author'), author)
    del_element.set(qn('w:date'), date)

    # 创建 run 元素
    del_run = OxmlElement('w:r')

    # 复制格式（如果有）
    if rpr_element is not None:
        del_run.append(rpr_element)

    # 创建 delText 元素
    del_text = OxmlElement('w:delText')
    del_text.text = text_content
    del_text.set(qn('xml:space'), 'preserve')

    del_run.append(del_text)
    del_element.append(del_run)

    return del_element


def create_insert_element(author, date, text_content, revision_id, rpr_element=None):
    """创建插入标记元素"""
    ins_element = OxmlElement('w:ins')
    ins_element.set(qn('w:id'), str(revision_id))
    ins_element.set(qn('w:author'), author)
    ins_element.set(qn('w:date'), date)

    # 创建 run 元素
    ins_run = OxmlElement('w:r')

    # 复制格式（如果有）
    if rpr_element is not None:
        ins_run.append(rpr_element)

    # 创建 text 元素
    ins_text = OxmlElement('w:t')
    ins_text.text = text_content
    ins_text.set(qn('xml:space'), 'preserve')

    ins_run.append(ins_text)
    ins_element.append(ins_run)

    return ins_element


def find_row_by_segment_id(table, segment_id):
    """在表格中查找包含指定 segment_id 的行"""
    for row_idx, row in enumerate(table.rows):
        # 检查第一个单元格（Segment ID 列）
        first_cell_text = row.cells[0].text.strip()
        if first_cell_text == str(segment_id).strip():
            return row, row_idx
    return None, None


def replace_cell_with_track_changes(cell, old_text, new_text, author, date_str, revision_id):
    """
    使用追踪修订替换单元格内容

    Args:
        cell: 要修改的单元格
        old_text: 旧文本（用于验证）
        new_text: 新文本
        author: 作者名称
        date_str: 日期字符串
        revision_id: 修订 ID

    Returns:
        新的 revision_id
    """
    # 获取当前文本
    current_text = cell.text.strip()

    # 保护标签
    old_protected = protect_tags(old_text)
    new_protected = protect_tags(new_text)
    current_protected = protect_tags(current_text)

    # 验证文本（允许旧文本为空）
    allow_empty_old = (old_text is None) or (str(old_text).strip() == "")
    if not allow_empty_old:
        if old_protected not in current_protected:
            raise ValueError(
                f"文本不匹配 - 预期: '{old_text[:30]}...', "
                f"实际: '{current_text[:30]}...'"
            )

    # 获取第一个段落
    if len(cell.paragraphs) == 0:
        return revision_id

    paragraph = cell.paragraphs[0]
    p_element = paragraph._element

    # 保存第一个 run 的格式（如果有）
    rpr_element = None
    if paragraph.runs:
        first_run = paragraph.runs[0]._element
        rpr_nodes = first_run.findall(qn('w:rPr'))
        if rpr_nodes:
            # 复制格式元素
            rpr_element = rpr_nodes[0]

    # 清空段落内容
    for child in list(p_element):
        if child.tag in [qn('w:r'), qn('w:del'), qn('w:ins'),
                         qn('w:hyperlink'), qn('w:bookmarkStart'),
                         qn('w:bookmarkEnd')]:
            p_element.remove(child)

    # 添加删除标记（如果有旧文本）
    if old_protected.strip():
        del_element = create_delete_element(
            author, date_str, old_protected, revision_id, rpr_element
        )
        p_element.append(del_element)
        revision_id += 1

    # 添加插入标记（如果有新文本）
    if new_protected.strip():
        ins_element = create_insert_element(
            author, date_str, new_protected, revision_id, rpr_element
        )
        p_element.append(ins_element)
        revision_id += 1

    return revision_id


def update_translations(doc, translations, author='Claude'):
    """
    更新翻译

    Args:
        doc: Word 文档对象
        translations: 翻译列表
        author: 作者名称

    Returns:
        (successful_count, failed_list)
    """
    successful = 0
    failed = []

    # 获取主表格（假设是第一个表格）
    if not doc.tables:
        raise ValueError("文档中没有找到表格")

    table = doc.tables[0]
    date_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')
    revision_id = 0

    print(f"开始处理 {len(translations)} 个翻译...")
    print("=" * 80)

    for i, trans in enumerate(translations, 1):
        segment_id = trans.get('segment_id', 'Unknown')
        old_text = trans.get('old_text', '')
        new_text = trans.get('new_text', '')

        print(f"[{i}/{len(translations)}] 处理 {segment_id[:40]}...", end=' ')

        try:
            # 查找行
            row, row_idx = find_row_by_segment_id(table, segment_id)
            if row is None:
                raise ValueError(f"找不到 segment_id: {segment_id}")

            # 获取目标单元格（第 4 列，索引 3）
            if len(row.cells) < 4:
                raise ValueError("表格列数不足（需要至少 4 列）")

            target_cell = row.cells[3]

            # 应用追踪修订
            revision_id = replace_cell_with_track_changes(
                target_cell, old_text, new_text, author, date_str, revision_id
            )

            successful += 1
            print("✓")

        except Exception as e:
            error_msg = str(e)
            failed.append((segment_id, error_msg))
            print(f"✗ ({error_msg[:60]})")

    return successful, failed


def restore_tags_in_document(doc):
    """
    恢复文档中的标签保护
    将 ⟨51⟩ 还原为 <51>
    """
    print("\n恢复标签保护...")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            run.text = restore_tags(run.text)

    print("✓ 标签恢复完成")


def main():
    parser = argparse.ArgumentParser(
        description='FC Insider 翻译更新（简化版 - 直接操作 DOCX）',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  python update_fc_insider_simple.py \\
    --input input.docx \\
    --translations translations.json \\
    --output output.docx \\
    --author "Your Name"
        '''
    )

    parser.add_argument('--input', required=True, help='输入 DOCX 文件')
    parser.add_argument('--translations', required=True, help='翻译 JSON 文件路径')
    parser.add_argument('--output', required=True, help='输出 DOCX 文件')
    parser.add_argument('--author', default='Claude', help='作者名称（默认: Claude）')

    args = parser.parse_args()

    # 检查输入文件
    if not Path(args.input).exists():
        print(f"✗ 错误：输入文件不存在 - {args.input}")
        return 1

    if not Path(args.translations).exists():
        print(f"✗ 错误：翻译文件不存在 - {args.translations}")
        return 1

    try:
        # 加载翻译数据
        with open(args.translations, 'r', encoding='utf-8') as f:
            data = json.load(f)
            translations = data.get('translations', data)

        # 打开文档
        print(f"读取文档: {args.input}")
        doc = Document(args.input)

        # 启用追踪修订
        enable_track_changes(doc)

        print(f"FC Insider 翻译更新（简化版）")
        print(f"作者: {args.author}")
        print(f"翻译数量: {len(translations)}")
        print("=" * 80)

        # 更新翻译
        successful, failed = update_translations(doc, translations, args.author)

        # 恢复标签
        restore_tags_in_document(doc)

        print("=" * 80)
        print(f"✓ 更新完成: {successful}/{len(translations)}")

        if failed:
            print(f"✗ 失败: {len(failed)} 个")
            print("\n失敗詳情（前 10 個）:")
            for seg_id, error in failed[:10]:
                print(f"  • {seg_id[:40]}: {error[:200]}")

        print("=" * 80)

        # 保存文档
        print(f"\n保存文档: {args.output}")
        doc.save(args.output)
        print("✓ 文档已保存")

        return 0 if successful > 0 else 1

    except Exception as e:
        print(f"\n✗ 致命错误: {e}")
        import traceback
        traceback.print_exc()
        return 2


if __name__ == '__main__':
    import sys
    sys.exit(main())
